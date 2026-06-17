import io
import os
import re
import json
import base64
import datetime
import subprocess
import threading
import queue as _queue
import ollama
from PIL import ImageGrab
from ddgs import DDGS
import time

# keyboard is used for the Ctrl+Q abort shortcut.
# Install with:  pip install keyboard
# Note: on Windows, keyboard requires no extra drivers. Run as admin if hotkeys
# don't fire (rare — usually works in normal user sessions inside a terminal).
try:
    import keyboard as _keyboard
    _KEYBOARD_AVAILABLE = True
except ImportError:
    _KEYBOARD_AVAILABLE = False

# ── Optional file-format libraries ────────────────────────────────────────────
# PDF reading:   pip install pymupdf
# DOCX reading:  pip install mammoth
# DOCX writing:  pip install python-docx
try:
    import fitz as _fitz
    _PDF_AVAILABLE = True
except ImportError:
    _fitz = None
    _PDF_AVAILABLE = False

try:
    import mammoth as _mammoth
    _MAMMOTH_AVAILABLE = True
except ImportError:
    _mammoth = None
    _MAMMOTH_AVAILABLE = False

try:
    import docx as _docx
    _DOCX_AVAILABLE = True
except ImportError:
    _docx = None
    _DOCX_AVAILABLE = False


# rich renders Markdown in the terminal (headers, bold, code blocks, tables).
# Pure formatting — zero effect on model logic or performance.
# Install with:  pip install rich
try:
    from rich.console import Console as _Console
    from rich.markdown import Markdown as _Markdown
    _console          = _Console()
    _RICH_AVAILABLE   = True
except ImportError:
    _RICH_AVAILABLE   = False

def _print_reply(label: str, text: str):
    """Print Jarvis's reply, rendering Markdown if rich is available."""
    print(f"\n{label}")   # always print the label as plain text
    if _RICH_AVAILABLE and text.strip():
        _console.print(_Markdown(text))
    else:
        print(text)

# =============================================================================
# CONFIGURATION
# =============================================================================
STARTUP_DIR    = os.getcwd()
MODEL_NAME     = "jarvishehe"

# ── Legacy / weak native-tool-calling models ────────────────────────────────
# Ollama's `tools=` parameter relies on the model's chat template emitting a
# structured tool-call block that Ollama can parse into response["message"]["tool_calls"].
# Modern models (Llama 3.1+, Qwen3, etc., and "jarvis" built on them) do this
# reliably. Older/smaller models — qwen2.5-coder included — often *understand*
# tool-calling but emit the call as plain JSON (or <tool_call> tags) inside
# the regular `content` string instead of the structured field, so Ollama
# reports `tool_calls` as empty even though the model clearly tried.
#
# Models listed here (by substring match against MODEL_NAME) get a text-based
# fallback parser that scans `content` for a tool call when the native field
# comes back empty. Modern models are completely unaffected — this list is
# checked only as a fallback path that never runs unless tool_calls is empty.
LEGACY_TOOLCALL_MODELS = (
    "qwen2.5-coder",
    "qwen2.5",
    "codeqwen",
    "deepseek-coder",
    "codellama",
)

def _is_legacy_toolcall_model(model_name: str) -> bool:
    """True if model_name or its underlying base model matches a known weak tool-calling family."""
    low = model_name.lower()
    if any(fam in low for fam in LEGACY_TOOLCALL_MODELS):
        return True
    try:
        info = ollama.show(model_name)
        base = info.get("modelinfo", {}).get("general.basename", "").lower()
        return any(fam in base for fam in LEGACY_TOOLCALL_MODELS)
    except Exception:
        return False

STORAGE_DIR    = r"D:\Jarvis\jarvis_project\storage"
COMMANDS_FILE       = os.path.join(STORAGE_DIR, "commands.md")
INSTRUCTIONS_FILE   = os.path.join(STORAGE_DIR, "instructions.md")
PATHS_FILE          = os.path.join(STORAGE_DIR, "paths.md")
DOMAIN_INDEX        = os.path.join(STORAGE_DIR, "domain_index.md")
SKILLS_DIR          = os.path.join(STORAGE_DIR, "skills")
SKILLS_INDEX        = r"D:\Jarvis\jarvis_project\skills.md"
DOMAIN_SKILLS_INDEX = os.path.join(STORAGE_DIR, "domain_skills_index.md")
MASTER_MEMORY   = os.path.join(STORAGE_DIR, "master_memory.md")
SESSION_MEMORY  = os.path.join(STORAGE_DIR, "session_memory.md")
RESPONSE_MEMORY = os.path.join(STORAGE_DIR, "response_memory.md")
TARGET_DIR     = r"D:\Jarvis"
LOG_FILE       = os.path.join(TARGET_DIR, "chat_log.md")

# Path to secrets file — stored in a completely separate folder from the Jarvis
# project so uploading or publishing the entire Jarvis folder is safe.
#
# Default location: C:\Users\<YourUsername>\AppData\Local\JarvisSecrets\jarvis_secrets.json
# This folder is outside any project directory and is never touched by git or file uploads.
#
# To use a different location, change SECRETS_FILE to any absolute path you prefer,
# as long as it is outside the Jarvis folder, e.g.:
#   SECRETS_FILE = r"C:\Users\YourName\Documents\jarvis_secrets.json"
#
# Format of the file:  { "GEMINI_API_KEY": "your_key_here" }
#
SECRETS_FILE = os.path.join(
    os.path.expanduser("~"),          # C:\Users\<YourUsername>
    "AppData", "Local",
    "JarvisSecrets",                  # dedicated folder, nothing else goes here
    "jarvis_secrets.json"
)

GOAL_SECTION_HEADER = "## Current Goal"
GOAL_SECTION_END    = "## Goal History"

# Physical screen resolution
SCREEN_W = 2560
SCREEN_H = 1600

# The model's internal canvas — Ollama vision models downscale to 1024px long edge
MODEL_CANVAS_W = 1024
MODEL_CANVAS_H = int(1024 * SCREEN_H / SCREEN_W)   # = 640 for 2560x1600

# Scale factors: real_px = canvas_px * SCALE
SCALE_X = SCREEN_W / MODEL_CANVAS_W   # 2.5
SCALE_Y = SCREEN_H / MODEL_CANVAS_H   # 2.5

# Grid drawn on screenshots at canvas resolution; every GRID_STEP canvas-px
GRID_STEP = 100

# =============================================================================
# ABORT FLAG — Ctrl+Q sets this to stop the current response
# =============================================================================
# A threading.Event that process_chat_turn checks at every loop iteration.
# When set, the turn is abandoned and control returns to the input prompt.
_abort_event = threading.Event()

# =============================================================================
# TESSERACT SETUP
# =============================================================================
#
# INSTALLATION (Windows) — do this once before running Jarvis:
#
#   1. Download the installer from:
#      https://github.com/UB-Mannheim/tesseract/wiki
#      (choose "tesseract-ocr-w64-setup-*.exe" for 64-bit)
#
#   2. Run the installer. Default path: C:\Program Files\Tesseract-OCR\tesseract.exe
#      Make sure "Add to PATH" is checked, OR set TESSERACT_PATH below manually.
#
#   3. Install the Python binding:
#      pip install pytesseract
#
# If Tesseract is not installed, fallback_find_text() returns an error string
# and Jarvis falls back to grid-based fallback_click_grid automatically.
#
TESSERACT_PATH = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
_TESSERACT_AVAILABLE = False

try:
    import pytesseract
    pytesseract.pytesseract.tesseract_cmd = TESSERACT_PATH
    # Quick smoke-test without a real image
    pytesseract.get_tesseract_version()
    _TESSERACT_AVAILABLE = True
except Exception:
    pass   # Jarvis will warn at startup; fallback_find_text() will return a clear error


# ── UI Automation (UIA) SETUP ───────────────────────────────────────────────
# pip install uiautomation pywin32
try:
    import win32gui
    import uiautomation as auto
    _UIA_AVAILABLE = True
except ImportError:
    _UIA_AVAILABLE = False
    win32gui = None
    auto = None

# Create the global UIA navigator instance right below your other globals
if _UIA_AVAILABLE:
    # Control types that can host meaningful child content — used as scan roots
    _CONTAINER_TYPES = {
        "PaneControl", "GroupControl", "CustomControl", "ToolbarControl",
        "DocumentControl", "WindowControl", "TabControl", "TabItemControl",
        "TreeControl", "ScrollBarControl", "MenuBarControl", "MenuControl",
    }
    # Control types that represent something a user can interact with
    _ACTIONABLE_TYPES = {
        "ButtonControl", "EditControl", "CheckBoxControl", "HyperlinkControl",
        "ListItemControl", "MenuItemControl", "TabItemControl", "RadioButtonControl",
        "ComboBoxControl", "SliderControl", "ImageControl", "SplitButtonControl",
        "TextControl",  # many Electron icon-buttons report as plain Text/Image with a Name
    }

    class AppMapNavigator:
        """
        UIA navigator tuned for deep Electron/Chromium UI trees (VS Code, Modrinth,
        Discord, etc). These apps wrap real content in many layers of generic
        Pane/Group/Custom containers, often 8-15 levels deep, so shallow depth
        cutoffs (3-4) never reach anything actionable.

        Strategy:
          - discover_ui_subtrees: walk up to MAX_DISCOVER_DEPTH levels, collect
            ANY named/automation-id'd container as a candidate subtree, not just
            top-level ones. Electron apps nest meaningfully-named panes deep.
          - inspect_subtree_controls: walk up to MAX_INSPECT_DEPTH levels below
            a chosen subtree root, collecting all actionable controls regardless
            of how deep they sit.
          - If a subtree yields nothing, automatically retry one level shallower
            in the tree (parent) so Jarvis doesn't have to manually backtrack.
        """

        MAX_DISCOVER_DEPTH = 12
        MAX_INSPECT_DEPTH  = 12
        MAX_RESULTS        = 60   # cap result lists so they don't blow the context

        def __init__(self, maps_dir="storage/app_maps"):
            self.maps_dir = os.path.abspath(maps_dir)
            os.makedirs(self.maps_dir, exist_ok=True)
            self._live_cache = {}

        def _get_map_path(self, window_title: str) -> str:
            safe_name = "".join([c if c.isalnum() else "_" for c in window_title.lower()]).strip("_")
            return os.path.join(self.maps_dir, f"{safe_name}.json")

        def load_app_blueprint(self, window_title: str) -> dict:
            path = self._get_map_path(window_title)
            if os.path.exists(path):
                try:
                    with open(path, "r", encoding="utf-8") as f:
                        return json.load(f)
                except Exception:
                    pass
            return {"window_title": window_title, "subtrees": {}, "known_controls": {}}

        def save_app_blueprint(self, window_title: str, blueprint: dict):
            path = self._get_map_path(window_title)
            try:
                with open(path, "w", encoding="utf-8") as f:
                    json.dump(blueprint, f, indent=2)
            except Exception:
                pass

        def _find_window(self, window_title: str):
            """
            Find a window by exact title, falling back to substring match
            (case-insensitive) since Electron apps often have dynamic titles
            like 'main.py - Jarvis - Visual Studio Code'.

            If multiple windows match the substring (e.g. several VS Code
            windows or several Chrome windows are open), prefer:
              1. The current foreground window, if it's among the matches.
              2. Otherwise the match with the longest title (often the most
                 "complete"/active one — e.g. a window showing a file path
                 vs a bare "Visual Studio Code" placeholder window).
            This makes window selection deterministic instead of picking
            whatever EnumWindows happens to return first.
            """
            hwnd = win32gui.FindWindow(None, window_title)
            if hwnd:
                return hwnd

            # Substring fallback — enumerate all visible top-level windows
            matches = []
            def cb(h, _):
                if win32gui.IsWindowVisible(h) and win32gui.GetWindowText(h):
                    title = win32gui.GetWindowText(h)
                    if window_title.lower() in title.lower():
                        # Skip zero-size / minimized-to-nothing helper windows
                        try:
                            rect = win32gui.GetWindowRect(h)
                            if (rect[2] - rect[0]) <= 0 or (rect[3] - rect[1]) <= 0:
                                return
                        except Exception:
                            pass
                        matches.append((h, title))
            win32gui.EnumWindows(cb, None)

            if not matches:
                return None
            if len(matches) == 1:
                return matches[0][0]

            # Multiple matches — prefer the foreground window if it's one of them
            try:
                fg = win32gui.GetForegroundWindow()
                for h, _title in matches:
                    if h == fg:
                        return h
            except Exception:
                pass

            # Otherwise pick the one with the longest title (most specific)
            matches.sort(key=lambda m: len(m[1]), reverse=True)
            return matches[0][0]

        def discover_ui_subtrees(self, window_title: str) -> list[dict]:
            hwnd = self._find_window(window_title)
            if not hwnd:
                return []

            root_element = auto.ControlFromHandle(hwnd)
            if window_title not in self._live_cache:
                self._live_cache[window_title] = {}

            containers = []
            seen_keys  = set()

            for element, depth in auto.WalkControl(root_element, maxDepth=self.MAX_DISCOVER_DEPTH):
                try:
                    type_name = element.ControlTypeName
                    name      = (element.Name or "").strip()
                    auto_id   = (element.AutomationId or "").strip()
                except Exception:
                    continue

                if type_name not in _CONTAINER_TYPES:
                    continue
                if not (name or auto_id):
                    continue

                key = f"{name or auto_id}::{type_name}::{depth}"
                if key in seen_keys:
                    continue
                seen_keys.add(key)

                self._live_cache[window_title][key] = element
                containers.append({
                    "subtree_key":  key,
                    "type":         type_name,
                    "name":         name,
                    "automation_id": auto_id,
                    "depth":        depth,
                })

                if len(containers) >= self.MAX_RESULTS:
                    break

            return containers

        def inspect_subtree_controls(self, window_title: str, subtree_key: str) -> list[dict]:
            app_cache = self._live_cache.get(window_title, {})
            subtree_root = app_cache.get(subtree_key)

            if not subtree_root:
                self.discover_ui_subtrees(window_title)
                subtree_root = self._live_cache.get(window_title, {}).get(subtree_key)
                if not subtree_root:
                    return []

            controls   = []
            seen_names = set()

            try:
                walker = auto.WalkControl(subtree_root, maxDepth=self.MAX_INSPECT_DEPTH)
            except Exception:
                return []

            for element, depth in walker:
                try:
                    type_name = element.ControlTypeName
                    name      = (element.Name or "").strip()
                    auto_id   = (element.AutomationId or "").strip()
                except Exception:
                    continue

                if type_name not in _ACTIONABLE_TYPES:
                    continue
                # TextControl is extremely common as static labels — only keep
                # it if it has an AutomationId (suggests it's a real interactive
                # element wrapper, common in Electron) or looks like a button label
                if type_name == "TextControl" and not auto_id:
                    continue
                if not (name or auto_id):
                    continue

                dedup_key = f"{name}::{auto_id}::{type_name}"
                if dedup_key in seen_names:
                    continue
                seen_names.add(dedup_key)

                controls.append({
                    "type": type_name, "name": name, "automation_id": auto_id, "depth": depth
                })

                if len(controls) >= self.MAX_RESULTS:
                    break

            return controls

        def safely_trigger_ui_element(self, window_title: str, control_type: str,
                                       search_property: str, property_value: str,
                                       action: str, text_to_type: str = "") -> str:
            hwnd = self._find_window(window_title)
            if not hwnd:
                return f"Error: Window '{window_title}' not found."

            root = auto.ControlFromHandle(hwnd)
            search_args = {}
            if search_property == "automation_id":
                search_args["AutomationId"] = property_value
            elif search_property == "name":
                search_args["Name"] = property_value
            elif search_property == "class_name":
                search_args["ClassName"] = property_value

            c_type = control_type.lower()
            # searchDepth=0 means unlimited in uiautomation — search the whole subtree
            search_args["searchDepth"] = 0

            if c_type == "button":
                element = root.ButtonControl(**search_args)
            elif c_type in ("edit", "input"):
                element = root.EditControl(**search_args)
            elif c_type == "text":
                element = root.TextControl(**search_args)
            elif c_type == "checkbox":
                element = root.CheckBoxControl(**search_args)
            elif c_type == "menuitem":
                element = root.MenuItemControl(**search_args)
            elif c_type == "listitem":
                element = root.ListItemControl(**search_args)
            elif c_type == "image":
                element = root.ImageControl(**search_args)
            else:
                element = root.Control(**search_args)

            if not element.Exists(1, 0.5):
                return (
                    f"Error: Could not locate {control_type} matching "
                    f"{search_property}='{property_value}' anywhere in '{window_title}'. "
                    f"Try manual_inspect_app_subtree on a different container, or fall back to "
                    f"fallback_click_text if this element has visible text."
                )

            try:
                if action == "click":
                    try:
                        element.GetInvokePattern().Invoke()
                        return f"Programmatically invoked '{property_value}'."
                    except Exception:
                        try:
                            element.SetFocus()
                        except Exception:
                            pass
                        element.Click(simulateMove=False)
                        return f"Background clicked '{property_value}'."
                elif action == "set_text":
                    try:
                        element.GetValuePattern().SetValue(text_to_type)
                    except Exception:
                        element.SetValue(text_to_type)
                    return f"Populated text field '{property_value}'."
                elif action == "get_text":
                    try:
                        return element.GetValuePattern().Value
                    except Exception:
                        return element.Name
            except Exception as e:
                return f"Action failed: {str(e)}"
            return "Unknown action."

        def find_and_act(self, window_title: str, description: str, action: str = "click",
                          text_to_type: str = "") -> str:
            """
            ONE-CALL UI interaction. Replaces the discover -> inspect -> interact chain.

            1. Find the window (fuzzy title match).
            2. Walk the ENTIRE control tree once (depth up to MAX_INSPECT_DEPTH),
               collecting every actionable element with its Name, AutomationId,
               ControlType, and screen bounding rectangle.
            3. Score each element against `description` using simple token-overlap
               + substring matching on Name and AutomationId. Pick the best match.
            4. Try to act on it three ways, in order:
               a. UIA InvokePattern.Invoke() — cleanest, works for native controls
               b. UIA Click() — simulated click via UIA, still works for some Electron
               c. Coordinate click via _do_click() at the element's screen-rect
                  center — this is the universal fallback. UIA can almost always
                  REPORT an element's position even when it can't invoke it
                  (this is exactly the Claude/Electron close-button case).
            For set_text / get_text, only (a)/direct pattern access is attempted —
            coordinate fallback doesn't apply to those.

            Returns a human-readable result string.
            """
            hwnd = self._find_window(window_title)
            if not hwnd:
                return (
                    f"Error: Window '{window_title}' not found. "
                    f"Call list_active_windows to see exact titles."
                )

            root = auto.ControlFromHandle(hwnd)

            # ── Walk the whole tree once, collect candidates ──────────────────────
            candidates = []
            try:
                for element, depth in auto.WalkControl(root, maxDepth=self.MAX_INSPECT_DEPTH):
                    try:
                        type_name = element.ControlTypeName
                        name      = (element.Name or "").strip()
                        auto_id   = (element.AutomationId or "").strip()
                    except Exception:
                        continue

                    if type_name not in _ACTIONABLE_TYPES and type_name not in _CONTAINER_TYPES:
                        continue
                    if not (name or auto_id):
                        continue

                    candidates.append({
                        "element": element, "type": type_name,
                        "name": name, "automation_id": auto_id, "depth": depth
                    })

                    if len(candidates) >= 400:   # hard safety cap on tree walk
                        break
            except Exception as e:
                return f"Error walking UI tree for '{window_title}': {str(e)}"

            if not candidates:
                return (
                    f"No interactive elements found in '{window_title}' at all. "
                    f"This window likely renders via canvas/WebGL (UIA cannot see "
                    f"inside it). Use fallback_click_text instead."
                )

            # ── Score candidates against the description ──────────────────────────
            desc_lower  = description.strip().lower()
            desc_tokens = set(re.findall(r"[a-z0-9]+", desc_lower))

            def score(c):
                name_l = c["name"].lower()
                id_l   = c["automation_id"].lower()
                s = 0.0

                # Exact match is by far the strongest signal
                if desc_lower == name_l:
                    s += 200
                elif desc_lower == id_l:
                    s += 150  # exact AutomationId match — slightly less certain than Name

                # Substring containment — only count for non-trivial descriptions
                # to avoid e.g. "ok" matching everything containing "ok"
                if len(desc_lower) >= 3:
                    if desc_lower in name_l:
                        s += 40
                    if desc_lower in id_l:
                        s += 20

                name_tokens = set(re.findall(r"[a-z0-9]+", name_l))
                id_tokens   = set(re.findall(r"[a-z0-9]+", id_l))
                s += 8 * len(desc_tokens & name_tokens)
                s += 3 * len(desc_tokens & id_tokens)

                # STRONGLY prefer actionable controls over containers — a Pane
                # or Group with a matching name is almost never the click target
                if c["type"] in _ACTIONABLE_TYPES:
                    s += 30
                else:
                    s -= 20

                # Slight preference for shallower (more likely top-level / visible) matches
                s -= c["depth"] * 0.2
                return s

            scored = sorted(candidates, key=score, reverse=True)
            best       = scored[0]
            best_score = score(best)

            if best_score <= 0:
                # No reasonable match — show top names so the model can retry with better wording
                actionable_sample = [c for c in candidates if c["type"] in _ACTIONABLE_TYPES][:25]
                sample_list = actionable_sample if actionable_sample else candidates[:25]
                sample = ", ".join(
                    f"'{c['name'] or c['automation_id']}' ({c['type']})"
                    for c in sample_list
                )
                return (
                    f"No element matched '{description}' well enough in '{window_title}'. "
                    f"Available elements include: {sample}{'...' if len(candidates) > len(sample_list) else ''}. "
                    f"Try again with one of these exact names, or use fallback_click_text."
                )

            # ── Disambiguation: if multiple elements scored very close to the best, ─
            # warn rather than silently picking one — picking the wrong one of
            # several near-identical matches (e.g. multiple "OK" buttons in
            # different dialogs) is a common source of unreliable clicks.
            close_matches = [c for c in scored[1:6] if best_score - score(c) <= 5]
            if close_matches and best_score < 100:
                # Only warn when the match isn't a clean exact-name match —
                # exact matches (score >= 200) are trusted even if duplicates exist
                # (clicking the first instance is usually correct, e.g. toolbar
                # buttons that legitimately appear in multiple panels).
                ambiguous_names = ", ".join(
                    f"'{c['name'] or c['automation_id']}' ({c['type']}, depth {c['depth']})"
                    for c in [best] + close_matches[:4]
                )
                print(f"   [UIA] ⚠️  Multiple similar matches for '{description}': {ambiguous_names}")
                print(f"   [UIA] Proceeding with best match: '{best['name'] or best['automation_id']}'")

            element   = best["element"]
            elem_desc = best["name"] or best["automation_id"] or best["type"]

            # ── get_text / set_text: pattern-based only, no coordinate fallback ────
            if action == "get_text":
                try:
                    return f"Text of '{elem_desc}': {element.GetValuePattern().Value}"
                except Exception:
                    return f"Text of '{elem_desc}': {element.Name}"

            if action == "set_text":
                try:
                    try:
                        element.GetValuePattern().SetValue(text_to_type)
                    except Exception:
                        element.SetValue(text_to_type)
                    return f"Success: set text of '{elem_desc}' to '{text_to_type[:40]}'"
                except Exception as e:
                    try:
                        element.Click(simulateMove=False)
                        type_text(text_to_type)
                        return (
                            f"Set text via click+type fallback on '{elem_desc}' "
                            f"(direct SetValue failed: {e})"
                        )
                    except Exception as e2:
                        return f"Error setting text on '{elem_desc}': {e2}"

            # ── click: try Invoke -> Click -> coordinate fallback ──────────────────
            # (a) InvokePattern
            try:
                element.GetInvokePattern().Invoke()
                return f"Success: invoked '{elem_desc}' (UIA InvokePattern)."
            except Exception:
                pass

            # (b) UIA Click()
            try:
                try:
                    element.SetFocus()
                except Exception:
                    pass
                element.Click(simulateMove=False)
                return f"Success: clicked '{elem_desc}' (UIA Click)."
            except Exception:
                pass

            # (c) Coordinate fallback — UIA can report position even if it can't invoke
            try:
                rect = element.BoundingRectangle
                if rect.width() <= 0 or rect.height() <= 0:
                    return (
                        f"Found '{elem_desc}' but it has no visible bounding box "
                        f"(width={rect.width()}, height={rect.height()}). "
                        f"It may be hidden or off-screen. Try fallback_click_text instead."
                    )
                cx = rect.left + rect.width()  // 2
                cy = rect.top  + rect.height() // 2
                result = _do_click(cx, cy, "left_click", label=f"UIA-located '{elem_desc}'")
                return (
                    f"'{elem_desc}' could not be invoked via UIA patterns directly, "
                    f"so Python clicked its on-screen position at ({cx},{cy}) instead. {result}"
                )
            except Exception as e:
                return (
                    f"Found '{elem_desc}' but could not determine its screen position "
                    f"({e}). Try fallback_click_text instead."
                )
        def read_aggregated_text(self, window_title: str, container_key: str = None) -> str:
            """
            Aggregates hundreds of small TextControl siblings into readable paragraphs.
            If container_key is provided, it reads only that subtree. 
            Otherwise, it attempts to read the main document area.
            """
            hwnd = self._find_window(window_title)
            if not hwnd: return "Window not found."
            
            root = auto.ControlFromHandle(hwnd)
            # If a specific container was requested, use that as the root
            search_root = self._live_cache.get(window_title, {}).get(container_key, root)
            
            aggregated_lines = []
            last_y = -1
            current_line = []
            
            # Walk the subtree, collecting TextControls
            for element, depth in auto.WalkControl(search_root, maxDepth=self.MAX_INSPECT_DEPTH):
                if element.ControlTypeName == "TextControl":
                    text = (element.Name or "").strip()
                    if not text: continue
                    
                    try:
                        rect = element.BoundingRectangle
                        # Logic: If elements are on the same Y-level (roughly), 
                        # treat them as part of the same line/paragraph
                        if abs(rect.top - last_y) < 15:
                            current_line.append(text)
                        else:
                            if current_line:
                                aggregated_lines.append(" ".join(current_line))
                            current_line = [text]
                            last_y = rect.top
                    except:
                        current_line.append(text)
            
            if current_line:
                aggregated_lines.append(" ".join(current_line))
                
            return "\n".join(aggregated_lines)
        def query_gemini_app(self, prompt: str) -> str:
            """
            Ensures the Gemini desktop app is running, injects a prompt into the UIA 
            EditControl interface, and uses text block aggregation to extract the output.
            """
            if not _UIA_AVAILABLE:
                return "Error: UIA engine is completely offline."

            window_title = "Gemini"
            # Step 1: Ensure window is active or launch it
            # 1. Ensure Gemini is running
            hwnd = win32gui.FindWindow(None, window_title)
            if not hwnd:
                print("   [UIA Bridge] Launching Gemini PWA via Desktop Shortcut...")
                desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
                
                # Find the shortcut (case-insensitive search)
                target_lnk = None
                for file in os.listdir(desktop_path):
                    if file.lower().endswith(".lnk") and "gemini" in file.lower():
                        target_lnk = os.path.join(desktop_path, file)
                        break
                
                if target_lnk:
                    # Use os.startfile to trigger the LNK file as if the user double-clicked it
                    os.startfile(target_lnk)
                    
                    # Await stabilization
                    for _ in range(15):
                        time.sleep(1.0)
                        hwnd = win32gui.FindWindow(None, "Gemini")
                        if hwnd: break
                else:
                    return "Error: Could not find 'Gemini' shortcut on Desktop."

            if not hwnd:
                return "Error: Gemini process started but window title failed to register."

            # Focus the target window context
            try:
                win32gui.ShowWindow(hwnd, 9) # Restore if minimized
                win32gui.SetForegroundWindow(hwnd)
                time.sleep(0.5)
            except Exception as e:
                print(f"   [Warning] Window focus constraint: {e}")

            root = auto.ControlFromHandle(hwnd)
            
            # Step 2: Locate the input field
            # Standard prompt boxes evaluate as EditControls. We hunt for standard input fields.
            input_field = None
            for element, depth in auto.WalkControl(root):
                if depth > 6: continue
                if element.ControlTypeName == "EditControl":
                    input_field = element
                    break # Grab primary input node

            if not input_field:
                return "Error: UIA tree traversal failed to isolate the primary message input field."

            # Step 3: Inject payload and trigger transmission
            try:
                input_field.SetFocus()
                input_field.SetValue(prompt)
                time.sleep(0.2)
                # Dispatch submission signal
                auto.SendKey(auto.Keys.VK_RETURN)
            except Exception as e:
                return f"Error: Input injection phase encountered a fault: {str(e)}"

            # Step 4: Await generation completion loop
            print("   [UIA Bridge] Prompt submitted. Tracking response generation pipeline...")
            time.sleep(4.0) # Initial network transit buffer
            
            # Read state loop to verify stabilization
            last_text_hash = 0
            stable_turns = 0
            aggregated_response = ""

            # Ensure the aggregator knows which window to read from

            for _ in range(25): # 25-second maximum processing ceiling
                time.sleep(1.0)
                # Leverage our internal line aggregator directly
                current_text = self.read_aggregated_text(window_title)
                current_hash = hash(current_text)
                
                if current_hash == last_text_hash and len(current_text.strip()) > len(prompt):
                    stable_turns += 1
                    if stable_turns >= 2: # Confirmed delta freeze across iterations
                        aggregated_response = current_text
                        break
                else:
                    stable_turns = 0
                    last_text_hash = current_hash

            if not aggregated_response:
                aggregated_response = self.read_aggregated_text(window_title)

            # Strip the injected user prompt out from the front of the aggregated layout
            clean_output = aggregated_response.replace(prompt, "").strip()
            return clean_output if clean_output else "Error: Response captured but content data is empty."
        
        def manage_gemini_chat(self, action: str, chat_name: str = None) -> str:
            window_title = "Gemini"
            hwnd = win32gui.FindWindow(None, window_title)
            if not hwnd: return "Error: Gemini window not found."
            
            root = auto.ControlFromHandle(hwnd)
            
            if action == "new_chat":
                # Search for the "New chat" button/icon
                for element, _ in auto.WalkControl(root):
                    if element.Name and "New chat" in element.Name:
                        element.Click()
                        return "Started a new Gemini chat."
                return "Error: Could not find 'New chat' button."
                
            elif action == "open_recent" and chat_name:
                # Search for the chat in the sidebar list
                for element, _ in auto.WalkControl(root):
                    if element.Name and chat_name.lower() in element.Name.lower():
                        element.Click()
                        return f"Opened recent chat: {chat_name}."
                return f"Error: Could not find recent chat named '{chat_name}'."
                
            return "Invalid action."
    ui_navigator = AppMapNavigator()
else:
    ui_navigator = None

# =============================================================================
# GEMINI SETUP (UPDATED FOR GOOGLE-GENAI & FREE TIER PROTECTION)
# =============================================================================
#
# INSTALLATION:
#   pip install google-genai
#
# API KEY SETUP (one-time):
#   1. Get a free key at https://aistudio.google.com/app/apikey
#   2. Create the secrets file at the path shown in SECRETS_FILE above.
#      Example:  { "GEMINI_API_KEY": "AIza..." }
#
# CREDIT DEFENSE MECHANISM:
#   To prevent running out of requests or hitting tight Token Per Minute (TPM) 
#   limits immediately on the free tier, we default routing to gemini-2.0-flash 
#   and window large text contexts safely.
#
_GEMINI_AVAILABLE = False
_gemini_client    = None

def _load_gemini():
    """Load API key from secrets file and initialise the unified GenAI client."""
    global _GEMINI_AVAILABLE, _gemini_client
    try:
        from google import genai
        secrets_path = os.path.abspath(SECRETS_FILE)
        if not os.path.exists(secrets_path):
            return False, f"Secrets file not found: {secrets_path}"
        with open(secrets_path, "r", encoding="utf-8") as f:
            secrets = json.load(f)
        key = secrets.get("GEMINI_API_KEY", "").strip()
        if not key:
            return False, "GEMINI_API_KEY is empty in secrets file."
        
        # Initialize modern unified client
        _gemini_client = genai.Client(api_key=key)
        _GEMINI_AVAILABLE = True
        return True, "OK"
    except ImportError:
        return False, "google-genai not installed. Run: pip install google-genai"
    except Exception as e:
        return False, str(e)


# Attempt load at module import time; failure is non-fatal
_gemini_load_ok, _gemini_load_msg = _load_gemini()


# Free-tier defense allocation profiles
_GEMINI_MODELS = {
    "quick":    ("gemini-2.0-flash", "Fast, high-efficiency tier for quick operations."),
    "balanced": ("gemini-2.0-flash", "Standard multi-step reasoning tier."),
    "hard":     ("gemini-2.5-flash", "Enhanced reasoning tier for code logic and architecture review."),
    "expert":   ("gemini-2.5-flash", "Capped safely to 2.5-flash to protect Free Tier limits from 2.5-pro exhaustion."),
}

_TASK_TYPE_KEYWORDS = {
    "quick":    ["summarise", "summarize", "lookup", "define", "translate", "short", "quick", "simple", "what is", "who is"],
    "hard":     ["analyse", "analyze", "compare", "reason", "plan", "debug", "review", "explain", "complex", "multi-step", "architecture", "design", "evaluate", "research"],
    "expert":   ["hardest", "expert", "deep", "comprehensive", "full analysis", "detailed review", "best possible"],
}

def _pick_gemini_model(prompt: str) -> tuple[str, str]:
    """
    Choose an appropriate fallback model protecting free credits.
    """
    low = prompt.lower()
    for task_type in ("expert", "hard", "quick"):
        if any(kw in low for kw in _TASK_TYPE_KEYWORDS[task_type]):
            m = _GEMINI_MODELS[task_type]
            return m[0], m[1]
    return _GEMINI_MODELS["balanced"]


def consult_gemini(prompt, task_type="auto", context=""):
    """
    Send a prompt via google.genai with smart character context-windowing.
    """
    if not _GEMINI_AVAILABLE or _gemini_client is None:
        return (
            f"Gemini is not available: {_gemini_load_msg}. "
            "Check SECRETS_FILE path and ensure google-genai is installed."
        )
    try:
        if task_type == "auto" or task_type not in _GEMINI_MODELS:
            model_id, rationale = _pick_gemini_model(prompt)
        else:
            model_id  = _GEMINI_MODELS[task_type][0]
            rationale = _GEMINI_MODELS[task_type][1]

        print(f"   [Gemini] Model: {model_id} — {rationale}")

        # Protect Free Tier TPM: Keep runaway contextual strings capped
        if len(context) > 60000:
            context = context[:30000] + "\n... [TRUNCATED CONTEXT TO PROTECT FREE API TIER LIMITS] ...\n" + context[-30000:]

        full_prompt = prompt.strip()
        if context.strip():
            full_prompt = "[CONTEXT]\n" + context.strip() + "\n\n[TASK]\n" + full_prompt

        # Modern unified API request syntax
        response = _gemini_client.models.generate_content(
            model=model_id,
            contents=full_prompt,
        )
        result = response.text.strip()

        print(f"   [Gemini] Response received safely ({len(result)} chars)")
        return f"[Gemini/{model_id}]\n" + result
    except Exception as e:
        return f"Gemini API error via google.genai: {str(e)}"


# =============================================================================
# 1. TOOL SCHEMAS
# =============================================================================

tools = [
    {
        "type": "function",
        "function": {
            "name": "manual_scan_app_layouts",
            "description": "Scan an active window to find its major layout containers (subtrees). Use this first when exploring a new app.",
            "parameters": {
                "type": "object",
                "properties": {
                    "window_title": {"type": "string", "description": "Exact title of the window (e.g., 'Gemini')."}
                },
                "required": ["window_title"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "manual_inspect_app_subtree",
            "description": "Scan a specific layout container (found via manual_scan_app_layouts) to reveal the interactive buttons and text fields inside it.",
            "parameters": {
                "type": "object",
                "properties": {
                    "window_title": {"type": "string"},
                    "subtree_key": {"type": "string", "description": "The name or automation_id of the container."}
                },
                "required": ["window_title", "subtree_key"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "click_ui_element",
            "description": (
                "PREFERRED tool for interacting with desktop app UI elements (buttons, "
                "fields, menu items, close/minimize buttons, etc). ONE call does "
                "everything: finds the window, searches its ENTIRE UI tree for an "
                "element matching your plain-English description, and acts on it. "
                "If the element can't be invoked directly via UI Automation (common "
                "for Electron app window-control buttons like Close/Minimize), it "
                "automatically falls back to clicking the element's exact on-screen "
                "position — no coordinates needed from you. "
                "Use this INSTEAD of manual_scan_app_layouts + manual_inspect_app_subtree + "
                "manual_interact_with_ui — those are for manual exploration only. "
                "If this returns 'No element matched', it will show you a list of "
                "available element names — retry with one of those exact names. "
                "If it says the window renders via canvas/WebGL, use "
                "fallback_click_text instead."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "window_title": {
                        "type": "string",
                        "description": "Window title or substring of it, e.g. 'Claude', 'Visual Studio Code'."
                    },
                    "description": {
                        "type": "string",
                        "description": "Plain-English description of the element, e.g. 'Close button', 'Send message', 'Settings icon'."
                    },
                    "action": {
                        "type": "string",
                        "description": "'click' (default), 'set_text', or 'get_text'.",
                        "enum": ["click", "set_text", "get_text"]
                    },
                    "text_to_type": {
                        "type": "string",
                        "description": "Required if action is 'set_text' — the text to enter."
                    }
                },
                "required": ["window_title", "description"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "manual_interact_with_ui",
            "description": (
                "ADVANCED/MANUAL: precise control via exact automation_id/name/class_name. "
                "Most tasks should use click_ui_element instead — it's a single call and "
                "has automatic coordinate fallback. Only use this if click_ui_element "
                "failed and you have an EXACT automation_id from manual_inspect_app_subtree."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "window_title": {"type": "string"},
                    "control_type": {"type": "string", "description": "'button', 'edit', 'text', etc."},
                    "search_property": {"type": "string", "description": "'automation_id', 'name', or 'class_name'"},
                    "property_value": {"type": "string", "description": "The target identifier value."},
                    "action": {"type": "string", "description": "'click', 'set_text', or 'get_text'"},
                    "text_to_type": {"type": "string", "description": "Only required if action is 'set_text'."}
                },
                "required": ["window_title", "control_type", "search_property", "property_value", "action"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "query_gemini_app",
            "description": "Opens the Gemini desktop application if it is closed, sends a complex natural language prompt, waits for the response generation to conclude, and extracts the full consolidated textual response back to Jarvis.",
            "parameters": {
                "type": "object",
                "properties": {
                    "prompt": {
                        "type": "string",
                        "description": "The complex query, data processing prompt, or reasoning task to pass down to the Gemini application window."
                    }
                },
                "required": ["prompt"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "manage_gemini_chat",
            "description": "Manage the Gemini application state by performing actions like starting a new chat or selecting a recent chat.",
            "parameters": {
                "type": "object",
                "properties": {
                    "action": {
                        "type": "string",
                        "enum": ["new_chat", "open_recent"],
                        "description": "The action to perform."
                    },
                    "chat_name": {
                        "type": "string",
                        "description": "Required if action is 'open_recent'. The specific name of the chat to open."
                    }
                },
                "required": ["action"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "list_active_windows",
            "description": "List all currently open and visible window titles on the desktop. Use this if you are unsure of the exact window_title to pass to the UI interaction tools.",
            "parameters": {
                "type": "object",
                "properties": {}
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "read_aggregated_text",
            "description": "Read text from a window or specific container by merging sibling TextControl elements into readable paragraphs.",
            "parameters": {
                "type": "object",
                "properties": {
                    "window_title": {"type": "string"},
                    "container_key": {"type": "string", "description": "Optional: Specific subtree key to read from."}
                },
                "required": ["window_title"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "read_local_file",
            "description": "Read the contents of a local file.",
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Absolute or relative path to the file."}
                },
                "required": ["path"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "write_local_file",
            "description": "Completely overwrite or clear a file. Pass '' to wipe it clean.",
            "parameters": {
                "type": "object",
                "properties": {
                    "path":    {"type": "string"},
                    "content": {"type": "string"}
                },
                "required": ["path", "content"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "append_local_file",
            "description": "Add content to the end of a file without modifying existing content.",
            "parameters": {
                "type": "object",
                "properties": {
                    "path":    {"type": "string"},
                    "content": {"type": "string"}
                },
                "required": ["path", "content"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "execute_terminal_command",
            "description": "Run a PowerShell command on the local Windows machine.",
            "parameters": {
                "type": "object",
                "properties": {
                    "command":           {"type": "string", "description": "The exact PowerShell command to execute."},
                    "working_directory": {"type": "string", "description": "Optional: absolute folder path to run from."}
                },
                "required": ["command"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "search_internet",
            "description": "Search the internet for real-time information or documentation.",
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {"type": "string"}
                },
                "required": ["query"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "fallback_view_screen",
            "description": (
                "Capture a live screenshot of the desktop, downscaled to canvas size with a "
                "coordinate grid burned in. Use this when you need to visually inspect the screen "
                "or when you need to identify coordinates for fallback_click_grid. "
                "For text-heavy GUIs, prefer fallback_find_text over reading the grid manually."
            ),
            "parameters": {"type": "object", "properties": {}},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "fallback_find_text",
            "description": (
                "Use OCR (Tesseract) to locate a text string on the current screen. "
                "Returns the canvas coordinates of the best match and a full list of all detected "
                "text with their positions. Use this instead of reading the grid image when you "
                "want to click a button, label, or menu item that has visible text — it is faster "
                "and more accurate than visual grid estimation. "
                "Pass the exact text or a substring of it. Case-insensitive."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "text": {
                        "type": "string",
                        "description": "The text string to search for on screen. Case-insensitive substring match."
                    }
                },
                "required": ["text"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "fallback_click_grid",
            "description": (
                "Simulate a mouse click at canvas coordinates from the grid screenshot. "
                "Python scales these to real screen pixels automatically. "
                "Use this when you have read coordinates from the fallback_view_screen grid. "
                "For clicking text elements, prefer fallback_find_text which gives you "
                "precise coordinates without needing to read the grid."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "x": {"type": "integer", "description": "Canvas x-coordinate from the grid screenshot."},
                    "y": {"type": "integer", "description": "Canvas y-coordinate from the grid screenshot."},
                    "click_type": {
                        "type": "string",
                        "description": "'left_click' (default), 'right_click', or 'double_click'.",
                        "enum": ["left_click", "right_click", "double_click"]
                    }
                },
                "required": ["x", "y"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "fallback_click_text",
            "description": (
                "Find a text element on screen using OCR and click it in one step. "
                "This is the most accurate way to click buttons, menu items, and labels. "
                "Use this whenever the element you want to click has readable text. "
                "If multiple matches exist, clicks the one with the highest OCR confidence."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "text": {
                        "type": "string",
                        "description": "The visible text of the element to click. Case-insensitive substring match."
                    },
                    "click_type": {
                        "type": "string",
                        "description": "'left_click' (default), 'right_click', or 'double_click'.",
                        "enum": ["left_click", "right_click", "double_click"]
                    }
                },
                "required": ["text"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "type_text",
            "description": (
                "Type text at the current cursor position using keyboard simulation. "
                "Click the target field first. Use special_key for Enter, Tab, Escape, F-keys etc."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "text":        {"type": "string", "description": "The text to type."},
                    "special_key": {"type": "string", "description": "Optional key to press after typing: 'Enter', 'Tab', 'Escape', 'F5', etc."}
                },
                "required": ["text"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "update_memory",
            "description": "Persist important information. 'target' must be 'master', 'project', or 'session'.",
            "parameters": {
                "type": "object",
                "properties": {
                    "target":  {"type": "string"},
                    "content": {"type": "string"}
                },
                "required": ["target", "content"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "set_current_goal",
            "description": "Update the current goal. Use goal='none' to clear when a task is done.",
            "parameters": {
                "type": "object",
                "properties": {
                    "goal":   {"type": "string"},
                    "reason": {"type": "string"}
                },
                "required": ["goal"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "load_skill",
            "description": "Load a skill file by name (without .md). Call list_skills first if unsure.",
            "parameters": {
                "type": "object",
                "properties": {
                    "skill_name": {"type": "string"}
                },
                "required": ["skill_name"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "list_skills",
            "description": "List all available skills with descriptions.",
            "parameters": {"type": "object", "properties": {}},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "read_instructions",
            "description": (
                "Read instructions.md — user preferences and behavioural rules "
                "(e.g. preferred command style, formatting, workflow habits). "
                "Consult before any task where HOW matters, not just what to do."
            ),
            "parameters": {"type": "object", "properties": {}},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "add_instruction",
            "description": (
                "Add a preference or behavioural rule to instructions.md. "
                "Call when the user states a preference or corrects your behaviour."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "instruction": {"type": "string", "description": "The rule to record."}
                },
                "required": ["instruction"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "read_paths",
            "description": (
                "Read paths.md — absolute paths to apps, folders, and files. "
                "Consult when you need a path you are not certain of."
            ),
            "parameters": {"type": "object", "properties": {}},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "add_path",
            "description": "Add a labelled path entry to paths.md.",
            "parameters": {
                "type": "object",
                "properties": {
                    "label": {"type": "string", "description": "Short name, e.g. 'Blender'."},
                    "path":  {"type": "string", "description": "Absolute path on disk."},
                    "note":  {"type": "string", "description": "Optional extra context."}
                },
                "required": ["label", "path"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "create_domain_knowledge",
            "description": (
                "Create a domain-specific knowledge file (like commands.md but for a specific "
                "tool, e.g. blender_commands.md). Registered in domain_index.md."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "name":            {"type": "string", "description": "snake_case name, no extension."},
                    "description":     {"type": "string", "description": "One-line description."},
                    "initial_content": {"type": "string", "description": "Optional seed content."}
                },
                "required": ["name", "description"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "list_domain_knowledge",
            "description": "List all registered domain knowledge files with descriptions.",
            "parameters": {"type": "object", "properties": {}},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "read_domain_knowledge",
            "description": "Read a domain knowledge file by name (without .md extension).",
            "parameters": {
                "type": "object",
                "properties": {
                    "name": {"type": "string", "description": "Filename without extension."}
                },
                "required": ["name"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "create_domain_skill",
            "description": (
                "Create a domain-specific skill file for a tool or workflow. "
                "Stored in skills dir, registered in domain_skills_index.md and skills.md."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "name":        {"type": "string", "description": "snake_case filename, no extension."},
                    "domain":      {"type": "string", "description": "Tool this belongs to, e.g. 'blender'."},
                    "description": {"type": "string", "description": "One-line description."},
                    "content":     {"type": "string", "description": "Full Markdown skill instructions."}
                },
                "required": ["name", "domain", "description", "content"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "list_domain_skills",
            "description": "List all domain-specific skills grouped by domain.",
            "parameters": {"type": "object", "properties": {}},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "consult_gemini",
            "description": (
                "Send a reasoning, analysis, or research task to a Gemini model via Google AI Studio. "
                "Use this when: (1) the user explicitly says 'ask Gemini' or 'consult Gemini', "
                "(2) the task requires deep reasoning, complex analysis, code review, architectural "
                "decisions, or multi-step planning that exceeds your own confident ability, "
                "(3) you need a second opinion or want to cross-check your own reasoning. "
                "Jarvis selects the most appropriate model automatically based on task complexity "
                "unless you specify task_type. "
                "Models available (free tier): "
                "quick=gemini-2.0-flash-lite (fast, simple tasks), "
                "balanced=gemini-2.0-flash (default, multi-step reasoning), "
                "hard=gemini-2.5-flash (complex analysis, long context), "
                "expert=gemini-2.5-pro (hardest problems, slowest)."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "prompt": {
                        "type": "string",
                        "description": "The full question or task for Gemini. Be specific and complete."
                    },
                    "task_type": {
                        "type": "string",
                        "description": "Model tier: 'auto' (default), 'quick', 'balanced', 'hard', or 'expert'."
                    },
                    "context": {
                        "type": "string",
                        "description": "Optional: relevant context to include (file contents, memory, prior results)."
                    }
                },
                "required": ["prompt"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "wait",
            "description": "Pause execution for a specific number of seconds. Use this when waiting for an application to launch, a web page to load, or a background process to complete.",
            "parameters": {
                "type": "object",
                "properties": {
                    "seconds": {
                        "type": "number",
                        "description": "The number of seconds to pause (can be a decimal, e.g., 1.5)."
                    }
                },
                "required": ["seconds"]
            }
        }
    },
    {"type":"function","function":{"name":"read_file_smart","description":"Read any file: txt/md/py/json/csv/html/.pdf(requires pymupdf)/.docx(requires mammoth). Returns chunk 1 for large files with chunk count — call read_file_chunk for rest.","parameters":{"type":"object","properties":{"path":{"type":"string","description":"Absolute path to the file."}},"required":["path"]}}},
    {"type":"function","function":{"name":"read_file_chunk","description":"Read chunk N (1-based) of a large file after read_file_smart reports multiple chunks.","parameters":{"type":"object","properties":{"path":{"type":"string"},"chunk_index":{"type":"integer","description":"1-based chunk number."}},"required":["path","chunk_index"]}}},
    {"type":"function","function":{"name":"write_docx_file","description":"Write a .docx Word document from Markdown-style text (# headings, **bold**). Requires python-docx: pip install python-docx.","parameters":{"type":"object","properties":{"path":{"type":"string","description":"Absolute path ending in .docx."},"content":{"type":"string","description":"Markdown-style text content."}},"required":["path","content"]}}},
    {"type":"function","function":{"name":"write_response_memory","description":"Overwrite the response scratchpad (response_memory.md). Call FIRST for any multi-step task with a numbered plan. Wiped automatically when set_current_goal(none) fires.","parameters":{"type":"object","properties":{"content":{"type":"string","description":"Plan, checklist, or notes."}},"required":["content"]}}},
    {"type":"function","function":{"name":"append_response_memory","description":"Append a note or partial result to the response scratchpad. Use to log progress and accumulate partial outputs during a task.","parameters":{"type":"object","properties":{"content":{"type":"string","description":"Note or partial result."}},"required":["content"]}}},
    {"type":"function","function":{"name":"read_response_memory","description":"Read the current response scratchpad to check your plan or assemble a final answer from accumulated notes.","parameters":{"type":"object","properties":{}}}},
]


# =============================================================================
# 2. LOCAL PYTHON FUNCTIONS
# =============================================================================

def search_internet(query):
    try:
        print(f" -> Searching the web for: '{query}'")
        with DDGS() as ddgs:
            results = list(ddgs.text(query, max_results=5))
        if not results:
            return "No web results found."
        return "\n".join(
            f"[{i}] Title: {r.get('title')}\n    Link: {r.get('href')}\n    Snippet: {r.get('body')}"
            for i, r in enumerate(results, 1)
        )
    except Exception as e:
        return f"Error executing internet search: {str(e)}"


CHUNK_CHARS = 12000   # max chars per file chunk sent to the model (~3k tokens)

def _encode_text(text: str) -> str:
    """Base64-encode text for safe transmission to the model."""
    b64 = base64.b64encode(text.encode("utf-8")).decode("utf-8")
    return "[SYSTEM NOTICE: Base64-encoded. Decode internally.]\nBASE64_PAYLOAD:\n" + b64

def read_local_file(path):
    try:
        with open(path, "r", encoding="utf-8") as f:
            content = f.read()
        return _encode_text(content)
    except Exception as e:
        return f"Error reading file: {str(e)}"

def read_file_smart(path):
    """Read any supported format. Large files are chunked at CHUNK_CHARS chars."""
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == ".pdf":
            if not _PDF_AVAILABLE:
                return "PDF reading requires PyMuPDF: pip install pymupdf"
            doc   = _fitz.open(path)
            pages = [page.get_text() for page in doc]
            doc.close()
            text  = "\n\n".join(f"[Page {i+1}]\n{p}" for i, p in enumerate(pages))
        elif ext == ".docx":
            if not _MAMMOTH_AVAILABLE:
                return "DOCX reading requires mammoth: pip install mammoth"
            with open(path, "rb") as f:
                result = _mammoth.extract_raw_text(f)
            text = result.value
        else:
            with open(path, "r", encoding="utf-8", errors="replace") as f:
                text = f.read()

        if len(text) > CHUNK_CHARS:
            total  = (len(text) + CHUNK_CHARS - 1) // CHUNK_CHARS
            header = (f"[FILE: {os.path.basename(path)} — {total} chunks of ~{CHUNK_CHARS} chars. "
                      f"This is chunk 1/{total}. "
                      f"Call read_file_chunk(path, N) for chunks 2..{total}]\n\n")
            return _encode_text(header + text[:CHUNK_CHARS])
        return _encode_text(f"[FILE: {os.path.basename(path)}]\n\n{text}")
    except Exception as e:
        return f"Error reading file: {str(e)}"

def read_file_chunk(path, chunk_index: int):
    """Read a specific chunk (1-based) of a large file."""
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == ".pdf":
            if not _PDF_AVAILABLE:
                return "PDF reading requires PyMuPDF: pip install pymupdf"
            doc   = _fitz.open(path)
            pages = [page.get_text() for page in doc]
            doc.close()
            text  = "\n\n".join(f"[Page {i+1}]\n{p}" for i, p in enumerate(pages))
        elif ext == ".docx":
            if not _MAMMOTH_AVAILABLE:
                return "DOCX reading requires mammoth: pip install mammoth"
            with open(path, "rb") as f:
                result = _mammoth.extract_raw_text(f)
            text = result.value
        else:
            with open(path, "r", encoding="utf-8", errors="replace") as f:
                text = f.read()

        total = (len(text) + CHUNK_CHARS - 1) // CHUNK_CHARS
        if chunk_index < 1 or chunk_index > total:
            return f"Chunk {chunk_index} out of range (1-{total})."
        start  = (chunk_index - 1) * CHUNK_CHARS
        header = f"[FILE: {os.path.basename(path)} — chunk {chunk_index}/{total}]\n\n"
        return _encode_text(header + text[start:start + CHUNK_CHARS])
    except Exception as e:
        return f"Error reading chunk: {str(e)}"


def write_local_file(path, content):
    try:
        parent = os.path.dirname(path)
        if parent:
            os.makedirs(parent, exist_ok=True)
        with open(path, "w", encoding="utf-8") as f:
            f.write(content)
        return f"Success: wrote data to {path}"
    except Exception as e:
        return f"Error writing file: {str(e)}"


def append_local_file(path, content):
    try:
        exists = os.path.exists(path)
        with open(path, "a", encoding="utf-8") as f:
            if exists:
                f.write("\n")
            f.write(content)
        return f"Success: appended data to {path}"
    except Exception as e:
        return f"Error appending to file: {str(e)}"


def write_docx_file(path, content):
    """Write a .docx from Markdown-style text (# headings, **bold** runs)."""
    if not _DOCX_AVAILABLE:
        return "DOCX writing requires python-docx: pip install python-docx"
    try:
        parent = os.path.dirname(path)
        if parent:
            os.makedirs(parent, exist_ok=True)
        document = _docx.Document()
        for line in content.splitlines():
            s = line.rstrip()
            if s.startswith("### "):    document.add_heading(s[4:], level=3)
            elif s.startswith("## "):  document.add_heading(s[3:], level=2)
            elif s.startswith("# "):   document.add_heading(s[2:], level=1)
            elif s == "":              document.add_paragraph("")
            else:
                p     = document.add_paragraph()
                parts = re.split(r"(\*\*[^*]+\*\*)", s)
                for part in parts:
                    if part.startswith("**") and part.endswith("**"):
                        p.add_run(part[2:-2]).bold = True
                    else:
                        p.add_run(part)
        document.save(path)
        return f"Success: wrote DOCX to {path}"
    except Exception as e:
        return f"Error writing DOCX: {str(e)}"

def write_response_memory(content):
    """Overwrite the response scratchpad."""
    return write_local_file(RESPONSE_MEMORY, content)

def append_response_memory(content):
    """Append a note to the response scratchpad."""
    return append_local_file(RESPONSE_MEMORY, content)

def read_response_memory():
    """Read the current response scratchpad."""
    if not os.path.exists(RESPONSE_MEMORY) or os.path.getsize(RESPONSE_MEMORY) == 0:
        return "Response memory is empty."
    return read_local_file(RESPONSE_MEMORY)

def clear_response_memory():
    """Wipe the response scratchpad. Called automatically on set_current_goal(none)."""
    try:
        write_local_file(RESPONSE_MEMORY, "")
        print("\U0001f5d2  [Response memory cleared.]")
        return "Success: response memory cleared."
    except Exception as e:
        return f"Error clearing response memory: {str(e)}"

def explore_path(path):
    """List directory contents with plain Get-ChildItem (no extra flags)."""
    try:
        import platform
        if platform.system() == "Windows":
            ps = (
                f"Get-ChildItem -Path '{path}' | "
                f"Select-Object Name,"
                f"@{{n='Type';e={{if($_.PSIsContainer){{'Directory'}}else{{'File'}}}}}},"
                f"@{{n='Size';e={{if($_.PSIsContainer){{'-'}}else{{[math]::Round($_.Length/1KB,1).ToString()+'KB'}}}}}} | "
                f"Format-Table -AutoSize | Out-String"
            )
            result = execute_terminal_command(ps)
            stdout = result.split("STDOUT:")[-1].split("STDERR:")[0].strip()
            stderr = result.split("STDERR:")[-1].strip() if "STDERR:" in result else ""
            if stderr and "does not exist" in stderr.lower():
                return f"Path does not exist: {path}"
            return f"Contents of {path}:\n{stdout}" if stdout else f"Empty or not found: {path}"
        else:
            entries = os.listdir(path)
            lines   = [f"  {'DIR ' if os.path.isdir(os.path.join(path,e)) else 'FILE'}  {e}"
                       for e in sorted(entries)]
            return f"Contents of {path}:\n" + "\n".join(lines)
    except Exception as e:
        return f"Error exploring path: {str(e)}"

def execute_terminal_command(command, working_directory=None):
    try:
        cwd = working_directory if working_directory else STARTUP_DIR
        import platform
        if platform.system() == "Windows":
            result = subprocess.run(
                ["powershell", "-Command", command],
                capture_output=True, text=True, timeout=30, cwd=cwd
            )
        else:
            result = subprocess.run(
                command, shell=True, capture_output=True, text=True, timeout=30, cwd=cwd
            )
        return f"STDOUT:\n{result.stdout}\nSTDERR:\n{result.stderr}"
    except Exception as e:
        return f"Execution failed: {str(e)}"
    
def manual_scan_app_layouts(window_title: str):
    if not _UIA_AVAILABLE:
        return "UIA library not installed. pip install uiautomation pywin32"
    containers = ui_navigator.discover_ui_subtrees(window_title)
    if not containers:
        return (
            f"No named containers found in '{window_title}' within "
            f"{ui_navigator.MAX_DISCOVER_DEPTH} levels. The window may not "
            f"expose a UI Automation tree (common for games/canvas-based apps). "
            f"Try fallback_click_text or fallback_view_screen instead."
        )
    summary = (
        f"Found {len(containers)} container(s) in '{window_title}'. "
        f"Pick the most relevant subtree_key and call manual_inspect_app_subtree.\n"
    )
    return summary + json.dumps(containers, indent=2)

def manual_inspect_app_subtree(window_title: str, subtree_key: str):
    if not _UIA_AVAILABLE:
        return "UIA library not installed. pip install uiautomation pywin32"
    controls = ui_navigator.inspect_subtree_controls(window_title, subtree_key)
    if not controls:
        return (
            f"No actionable controls found in '{subtree_key}' within "
            f"{ui_navigator.MAX_INSPECT_DEPTH} levels below it. "
            f"Try a different subtree_key from manual_scan_app_layouts, or if this "
            f"app renders via canvas/WebGL (games, some web apps), use "
            f"fallback_click_text instead — UIA cannot see canvas content."
        )
    summary = f"Found {len(controls)} actionable control(s) in '{subtree_key}':\n"
    return summary + json.dumps(controls, indent=2)

def manual_interact_with_ui(window_title: str, control_type: str, search_property: str, property_value: str, action: str, text_to_type: str = ""):
    if not _UIA_AVAILABLE: return "UIA library not installed."
    return ui_navigator.safely_trigger_ui_element(
        window_title, control_type, search_property, property_value, action, text_to_type
    )

def click_ui_element(window_title: str, description: str, action: str = "click", text_to_type: str = ""):
    """
    ONE-CALL UI interaction: find an element in `window_title` matching
    `description` (plain English, e.g. "Close button", "Send message field")
    and act on it. Falls back to coordinate-click automatically if the
    element can't be invoked via UIA patterns directly (common for
    Electron/Chromium close/minimize/maximize buttons).
    """
    if not _UIA_AVAILABLE:
        return "UIA library not installed. pip install uiautomation pywin32"
    return ui_navigator.find_and_act(window_title, description, action, text_to_type)

def list_active_windows():
    """
    Returns a clean list of all visible, named windows currently open on the desktop.
    Filters out hidden processes and empty overlays.
    """
    if not _UIA_AVAILABLE: 
        return "UIA library not installed."
    
    def enum_win_callback(hwnd, window_list):
        # We only care about windows that are actually drawn on screen
        if win32gui.IsWindowVisible(hwnd):
            title = win32gui.GetWindowText(hwnd).strip()
            # Filter out empty strings and common invisible Windows overlays
            if title and title not in ["Program Manager", "Settings", "Microsoft Text Input Application"]:
                window_list.append(title)
    
    windows = []
    win32gui.EnumWindows(enum_win_callback, windows)
    
    # Remove duplicates and sort alphabetically
    unique_windows = sorted(list(set(windows)))
    
    if not unique_windows:
        return "No visible named windows found."
    
    return "Currently open windows:\n" + "\n".join(f"- {w}" for w in unique_windows)


# =============================================================================
# 3. SCREEN CAPTURE & OCR
# =============================================================================

def _grab_full_screenshot():
    """Grab the full-resolution screen and return a PIL Image."""
    return ImageGrab.grab()


def _scale_canvas_to_screen(cx, cy):
    """Convert canvas coordinates to real screen coordinates."""
    return int(round(cx * SCALE_X)), int(round(cy * SCALE_Y))


def _scale_screen_to_canvas(rx, ry):
    """Convert real screen coordinates to canvas coordinates."""
    return int(round(rx / SCALE_X)), int(round(ry / SCALE_Y))


def capture_screen_to_ram():
    """
    Grab screen → downscale to canvas → burn coordinate grid → return base64 JPEG.
    The grid labels are at canvas resolution. The model reads them and passes them
    directly to fallback_click_grid; Python scales back to real pixels.
    """
    try:
        from PIL import ImageDraw, ImageFont
        screenshot = _grab_full_screenshot()

        # Downscale to canvas
        from PIL import Image as _PILImage
        canvas = screenshot.resize((MODEL_CANVAS_W, MODEL_CANVAS_H), resample=_PILImage.LANCZOS)
        draw   = ImageDraw.Draw(canvas)
        cw, ch = canvas.size

        try:
            font = ImageFont.truetype("cour.ttf", 10)
        except Exception:
            font = ImageFont.load_default()

        line_col   = (60, 60, 60)
        label_fg   = (255, 255, 0)
        label_shad = (0, 0, 0)

        for x in range(0, cw, GRID_STEP):
            draw.line([(x, 0), (x, ch)], fill=line_col, width=1)
            draw.text((x + 2, 3), str(x), font=font, fill=label_shad)
            draw.text((x + 1, 2), str(x), font=font, fill=label_fg)

        for y in range(0, ch, GRID_STEP):
            draw.line([(0, y), (cw, y)], fill=line_col, width=1)
            draw.text((3, y + 2), str(y), font=font, fill=label_shad)
            draw.text((2, y + 1), str(y), font=font, fill=label_fg)

        buf = io.BytesIO()
        canvas.save(buf, format="JPEG", quality=75)
        return base64.b64encode(buf.getvalue()).decode("utf-8")
    except Exception as e:
        return f"Error capturing screen: {str(e)}"


# In-memory OCR cache: (screenshot_id, results_list)
# Avoids re-running Tesseract when find_text and click_text are called in the
# same turn from the same screenshot. Cache is invalidated by a new grab.
_ocr_cache: tuple = (None, None)   # (id(PIL_image), results)

def ocr_screen(screenshot=None):
    """
    Run Tesseract entirely in RAM — no temp files, no disk writes.

    pytesseract.image_to_data() accepts a PIL Image directly and pipes it
    to the tesseract process via stdin (using the 'pipe:' input method
    internally). No intermediate file is created on disk.

    Pass an existing PIL screenshot to reuse a grab; omit to grab fresh.
    Results are cached per PIL image object so the same screenshot is never
    OCR'd twice in one turn.

    Returns a list of word dicts or None if Tesseract is unavailable.
    """
    global _ocr_cache
    if not _TESSERACT_AVAILABLE:
        return None
    try:
        import pytesseract
        if screenshot is None:
            screenshot = _grab_full_screenshot()

        # Cache hit — same PIL object (same turn, same grab)
        if _ocr_cache[0] is id(screenshot):
            return _ocr_cache[1]

        # image_to_data with a PIL Image uses stdin piping internally —
        # no temp file is written to disk.
        data = pytesseract.image_to_data(
            screenshot,
            output_type=pytesseract.Output.DICT,
            nice=0,          # don't lower process priority
        )
        results = []
        n = len(data["text"])
        for i in range(n):
            word = data["text"][i].strip()
            conf = int(data["conf"][i])
            if not word or conf < 30:
                continue
            left = data["left"][i]
            top  = data["top"][i]
            w    = data["width"][i]
            h    = data["height"][i]
            sx   = left + w // 2
            sy   = top  + h // 2
            cx, cy = _scale_screen_to_canvas(sx, sy)
            results.append({
                "text":     word,
                "conf":     conf,
                "screen_x": sx,
                "screen_y": sy,
                "canvas_x": cx,
                "canvas_y": cy,
                "left": left, "top": top, "w": w, "h": h,
            })
        _ocr_cache = (id(screenshot), results)
        return results
    except Exception:
        return None


def fallback_find_text(text, _screenshot=None):
    """
    Tool implementation for fallback_find_text.
    Returns a structured text report of all matches with canvas coordinates.
    Pass _screenshot to reuse an existing grab (avoids a second screen capture).
    """
    if not _TESSERACT_AVAILABLE:
        return (
            "Tesseract OCR is not installed or not found. "
            "Cannot use text-based screen search. "
            "Install Tesseract from https://github.com/UB-Mannheim/tesseract/wiki "
            "and set TESSERACT_PATH in main.py. "
            "Fall back to fallback_view_screen + fallback_click_grid with grid coordinates."
        )
    words = ocr_screen(screenshot=_screenshot)
    if words is None:
        return "OCR failed — screen could not be read."

    query  = text.strip().lower()
    # Collect all words whose text contains the query (substring, case-insensitive)
    matches = [w for w in words if query in w["text"].lower()]

    if not matches:
        # Show everything Tesseract found so the model can adapt
        all_words = sorted(set(w["text"] for w in words))
        return (
            f"No text matching '{text}' found on screen.\n"
            f"All detected text on screen:\n"
            + ", ".join(f'"{w}"' for w in all_words[:80])
            + (" ... (truncated)" if len(all_words) > 80 else "")
        )

    # Sort by confidence descending; best match first
    matches.sort(key=lambda w: w["conf"], reverse=True)
    best = matches[0]

    lines = [
        f"Found {len(matches)} match(es) for '{text}'.",
        f"Best match: '{best['text']}' (conf={best['conf']}%) "
        f"at canvas ({best['canvas_x']}, {best['canvas_y']}) "
        f"→ screen ({best['screen_x']}, {best['screen_y']})",
        "",
        "All matches (canvas coords):",
    ]
    for m in matches[:10]:   # cap at 10 to keep output compact
        lines.append(
            f"  '{m['text']}' conf={m['conf']}% "
            f"canvas=({m['canvas_x']},{m['canvas_y']})"
        )
    return "\n".join(lines)


def fallback_click_text(text, click_type="left_click", _screenshot=None):
    """
    Find text on screen via OCR and click its center in one step.
    Pass _screenshot to reuse an existing grab.
    Returns a status string.
    """
    if not _TESSERACT_AVAILABLE:
        return (
            "Tesseract OCR is not installed. "
            "Use fallback_view_screen + fallback_click_grid instead."
        )
    words = ocr_screen(screenshot=_screenshot)
    if words is None:
        return "OCR failed — cannot locate text."

    query   = text.strip().lower()
    matches = [w for w in words if query in w["text"].lower()]
    if not matches:
        all_words = sorted(set(w["text"] for w in words))
        return (
            f"Text '{text}' not found on screen. "
            f"Detected text includes: {', '.join(repr(w) for w in all_words[:40])}"
        )

    matches.sort(key=lambda w: w["conf"], reverse=True)
    best = matches[0]
    sx, sy = best["screen_x"], best["screen_y"]
    cx, cy = best["canvas_x"], best["canvas_y"]

    print(f"   [OCR Click] '{best['text']}' conf={best['conf']}% "
          f"canvas({cx},{cy}) → screen({sx},{sy})")
    return _do_click(sx, sy, click_type, label=f"OCR '{best['text']}'")


def fallback_click_grid(x, y, click_type="left_click"):
    """
    x, y are CANVAS coordinates from the grid screenshot.
    Python scales to real screen pixels before clicking.
    """
    real_x, real_y = _scale_canvas_to_screen(x, y)
    print(f"   [Grid Click] canvas({x},{y}) → screen({real_x},{real_y})")
    return _do_click(real_x, real_y, click_type, label=f"grid ({x},{y})")


def _do_click(screen_x, screen_y, click_type, label=""):
    """
    Perform the actual click at real screen coordinates using PowerShell + user32.dll.
    Shared by both fallback_click_grid and fallback_click_text.
    """
    try:
        if click_type == "double_click":
            events = (
                "$m::mouse_event(0x0002,0,0,0,0)\n"
                "$m::mouse_event(0x0004,0,0,0,0)\n"
                "Start-Sleep -Milliseconds 50\n"
                "$m::mouse_event(0x0002,0,0,0,0)\n"
                "$m::mouse_event(0x0004,0,0,0,0)"
            )
        elif click_type == "right_click":
            events = (
                "$m::mouse_event(0x0008,0,0,0,0)\n"
                "$m::mouse_event(0x0010,0,0,0,0)"
            )
        else:
            events = (
                "$m::mouse_event(0x0002,0,0,0,0)\n"
                "$m::mouse_event(0x0004,0,0,0,0)"
            )

        ps_script = (
            f"Add-Type -AssemblyName System.Windows.Forms\n"
            f"[System.Windows.Forms.Cursor]::Position = "
            f"New-Object System.Drawing.Point({screen_x},{screen_y})\n"
            f"Start-Sleep -Milliseconds 50\n"
            f"$sig = '[DllImport(\"user32.dll\")] public static extern void "
            f"mouse_event(int flags, int dx, int dy, int data, int extra);'\n"
            f"$m = Add-Type -MemberDefinition $sig -Name 'Win32M' -Namespace W32 -PassThru\n"
            f"{events}"
        )
        result  = execute_terminal_command(ps_script)
        stderr  = result.split("STDERR:")[-1].strip() if "STDERR:" in result else ""
        if stderr:
            return f"{click_type} at screen({screen_x},{screen_y}) [{label}] — warning: {stderr[:150]}"
        return f"Success: {click_type} at screen({screen_x},{screen_y}) [{label}]"
    except Exception as e:
        return f"Error simulating click: {str(e)}"


def type_text(text, special_key=None):
    """Type text using PowerShell SendKeys at the current cursor position."""
    try:
        special_chars = "~%^+{}[]()"
        escaped = ""
        for ch in text:
            escaped += ("{" + ch + "}") if ch in special_chars else ch

        key_map = {
            "enter": "~", "tab": "{TAB}", "escape": "{ESC}",
            "backspace": "{BACKSPACE}", "delete": "{DELETE}",
            "home": "{HOME}", "end": "{END}",
            "pageup": "{PGUP}", "pagedown": "{PGDN}",
            "up": "{UP}", "down": "{DOWN}", "left": "{LEFT}", "right": "{RIGHT}",
        }
        if special_key:
            sk = special_key.lower()
            escaped += key_map.get(sk, "{" + special_key.upper() + "}")

        ps_script = (
            "Add-Type -AssemblyName System.Windows.Forms\n"
            f'[System.Windows.Forms.SendKeys]::SendWait("{escaped}")'
        )
        result = execute_terminal_command(ps_script)
        suffix = f" + {special_key}" if special_key else ""
        return f"Success: typed '{text[:40]}{'...' if len(text)>40 else ''}'{suffix}"
    except Exception as e:
        return f"Error typing text: {str(e)}"


# =============================================================================
# 4. KNOWLEDGE BASE — instructions.md, paths.md, domain files
# =============================================================================

def _ensure_kb_files():
    os.makedirs(STORAGE_DIR, exist_ok=True)
    if not os.path.exists(INSTRUCTIONS_FILE):
        write_local_file(INSTRUCTIONS_FILE,
            "# Jarvis Instructions & Preferences\n"
            "User preferences and behavioural rules.\n"
            "Format: one rule per line, starting with '- '.\n\n"
            "## Preferences\n")
    if not os.path.exists(PATHS_FILE):
        write_local_file(PATHS_FILE,
            "# Jarvis Paths\n"
            "Absolute paths to applications, folders and files.\n\n"
            "## Paths\n")
    if not os.path.exists(DOMAIN_INDEX):
        write_local_file(DOMAIN_INDEX,
            "# Jarvis Domain Knowledge Index\n"
            "Registered domain-specific knowledge files.\n"
            "Format: `filename_without_ext` - description\n\n"
            "## Files\n")
    if not os.path.exists(DOMAIN_SKILLS_INDEX):
        write_local_file(DOMAIN_SKILLS_INDEX,
            "# Jarvis Domain Skills Index\n"
            "Registered domain-specific skill files.\n"
            "Format: [domain] `filename_without_ext` - description\n\n"
            "## Skills\n")


def read_instructions():
    _ensure_kb_files()
    return read_local_file(INSTRUCTIONS_FILE)


def add_instruction(instruction):
    _ensure_kb_files()
    result = append_local_file(INSTRUCTIONS_FILE, f"- {instruction.strip()}")
    print(f"📌 [Instruction added]: {instruction.strip()[:80]}")
    return result


def read_paths():
    _ensure_kb_files()
    return read_local_file(PATHS_FILE)


def add_path(label, path, note=""):
    _ensure_kb_files()
    note_part = f"  _{note.strip()}_" if note.strip() else ""
    result = append_local_file(PATHS_FILE, f"- **{label.strip()}**: `{path.strip()}`{note_part}")
    print(f"📍 [Path added]: {label} -> {path}")
    return result


def create_domain_knowledge(name, description, initial_content=""):
    _ensure_kb_files()
    safe = re.sub(r'[^a-zA-Z0-9_]', '_', name.strip().lower())
    fpath = os.path.join(STORAGE_DIR, f"{safe}.md")
    if os.path.exists(fpath):
        return f"Domain knowledge '{safe}.md' already exists at {fpath}."
    header = (f"# Domain Knowledge: {safe}\n_{description.strip()}_\n\n"
              f"Created: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}\n\n")
    write_local_file(fpath, header + (initial_content.strip() + "\n" if initial_content.strip() else ""))
    append_local_file(DOMAIN_INDEX, f"- `{safe}` - {description.strip()}")
    print(f"📚 [Domain knowledge created]: {safe}.md")
    return f"Success: created '{safe}.md' at {fpath} and registered in domain index."


def list_domain_knowledge():
    _ensure_kb_files()
    return read_local_file(DOMAIN_INDEX)


def read_domain_knowledge(name):
    _ensure_kb_files()
    safe  = re.sub(r'[^a-zA-Z0-9_]', '_', name.strip().lower())
    fpath = os.path.join(STORAGE_DIR, f"{safe}.md")
    if not os.path.exists(fpath):
        try:
            match = next((e for e in os.listdir(STORAGE_DIR) if e.lower() == f"{safe}.md"), None)
            if match:
                fpath = os.path.join(STORAGE_DIR, match)
            else:
                return f"Domain knowledge '{safe}.md' not found. Call list_domain_knowledge."
        except Exception:
            return f"Domain knowledge '{safe}.md' not found."
    return read_local_file(fpath)


def create_domain_skill(name, domain, description, content):
    _ensure_kb_files()
    os.makedirs(SKILLS_DIR, exist_ok=True)
    safe  = re.sub(r'[^a-zA-Z0-9_]', '_', name.strip().lower())
    fpath = os.path.join(SKILLS_DIR, f"{safe}.md")
    if os.path.exists(fpath):
        return f"Domain skill '{safe}.md' already exists."
    header = (f"# Domain Skill: {safe}\n**Domain**: {domain.strip()}\n"
              f"_{description.strip()}_\n\n"
              f"Created: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}\n\n---\n\n")
    write_local_file(fpath, header + content.strip() + "\n")
    entry = f"- [{domain.strip()}] `{safe}` - {description.strip()}"
    append_local_file(DOMAIN_SKILLS_INDEX, entry)
    append_local_file(SKILLS_INDEX, entry)
    print(f"📋 [Domain skill created]: {safe}.md (domain: {domain})")
    return f"Success: created '{safe}.md' registered in both indexes."


def list_domain_skills():
    _ensure_kb_files()
    return read_local_file(DOMAIN_SKILLS_INDEX)


# =============================================================================
# 5. SKILL SYSTEM
# =============================================================================

def list_skills():
    os.makedirs(SKILLS_DIR, exist_ok=True)
    if not os.path.exists(SKILLS_INDEX):
        write_local_file(
            SKILLS_INDEX,
            "# Jarvis Skills Index\n\nSkill files live in: "
            + SKILLS_DIR + "\n\n## Skills\n_No skills registered yet._\n"
        )
        return "Skills index created. No skills registered yet."
    try:
        with open(SKILLS_INDEX, "r", encoding="utf-8") as f:
            content = f.read()
        b64 = base64.b64encode(content.encode("utf-8")).decode("utf-8")
        return (
            "[SYSTEM NOTICE: Base64-encoded. Decode internally.]\nBASE64_PAYLOAD:\n" + b64
        )
    except Exception as e:
        return f"Error reading skills index: {str(e)}"


def load_skill(skill_name):
    os.makedirs(SKILLS_DIR, exist_ok=True)
    skill_path = os.path.join(SKILLS_DIR, f"{skill_name}.md")
    if not os.path.exists(skill_path):
        try:
            entries = os.listdir(SKILLS_DIR)
            match = next(
                (e for e in entries if e.lower() == f"{skill_name.lower()}.md"), None
            )
            if match:
                skill_path = os.path.join(SKILLS_DIR, match)
            else:
                return f"Skill '{skill_name}' not found. Call list_skills to see available skills."
        except Exception:
            return f"Skill '{skill_name}' not found."
    try:
        with open(skill_path, "r", encoding="utf-8") as f:
            content = f.read()
        b64 = base64.b64encode(content.encode("utf-8")).decode("utf-8")
        print(f"📋 [Skill loaded: {skill_name}]")
        return (
            f"[SKILL LOADED: {skill_name}]\n"
            "[Decode internally and follow instructions exactly.]\n"
            "BASE64_PAYLOAD:\n" + b64
        )
    except Exception as e:
        return f"Error loading skill '{skill_name}': {str(e)}"


# =============================================================================
# 6. MEMORY SYSTEM
# =============================================================================

_active_project_memory_path = None
_current_goal               = None


def _memory_path_for_target(target):
    t = target.strip().lower()
    if t == "master":  return MASTER_MEMORY
    if t == "session": return SESSION_MEMORY
    if t == "project": return _active_project_memory_path
    return None


def update_memory(target, content):
    path = _memory_path_for_target(target)
    if not path:
        return f"Memory update skipped: no path for target '{target}'."
    ts    = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
    entry = f"[{ts}] {content.strip()}"
    result = append_local_file(path, entry)
    print(f"🧠 [Memory → {target}]: {content.strip()[:80]}{'...' if len(content)>80 else ''}")
    return result


def set_current_goal(goal, reason=""):
    global _current_goal
    ts         = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
    clean_goal = goal.strip()
    old_goal   = _current_goal or "none"

    raw = (
        open(SESSION_MEMORY, "r", encoding="utf-8").read()
        if os.path.exists(SESSION_MEMORY)
        else f"# Jarvis Session Memory\nSession started: {ts}\n"
    )
    lines      = raw.splitlines(keepends=True)
    header_end = 0
    for i, line in enumerate(lines):
        header_end = i
        if i > 0 and line.strip() == "":
            header_end = i + 1
            break

    header_block = "".join(lines[:header_end])
    body         = re.sub(
        r"## Current Goal.*?(?=\n## |\Z)", "", "".join(lines[header_end:]), flags=re.DOTALL
    ).lstrip("\n")

    goal_block  = (
        f"{GOAL_SECTION_HEADER}\n_No active goal._\n\n"
        if clean_goal.lower() == "none"
        else f"{GOAL_SECTION_HEADER}\n{clean_goal}\n\n"
    )
    new_content = header_block + goal_block + body
    reason_note = f" ({reason.strip()})" if reason.strip() else ""
    h_entry     = f"\n[{ts}] [GOAL CHANGED] {old_goal!r} → {clean_goal!r}{reason_note}"

    if GOAL_SECTION_END not in new_content:
        new_content += f"\n{GOAL_SECTION_END}\n{h_entry}\n"
    else:
        new_content = new_content.replace(GOAL_SECTION_END, GOAL_SECTION_END + h_entry)

    try:
        with open(SESSION_MEMORY, "w", encoding="utf-8") as f:
            f.write(new_content)
    except Exception as e:
        return f"Error writing session memory: {str(e)}"

    _current_goal = clean_goal if clean_goal.lower() != "none" else None
    label = f"Goal set: {clean_goal}" if _current_goal else "Goal cleared."
    print(f"🎯 [{label}]")
    if clean_goal.lower() == "none":
        clear_response_memory()
    return f"Success: {label}"


def get_current_goal_from_file():
    if not os.path.exists(SESSION_MEMORY):
        return None
    try:
        content = open(SESSION_MEMORY, "r", encoding="utf-8").read()
        m = re.search(r"## Current Goal\n(.+?)(?=\n## |\Z)", content, re.DOTALL)
        if not m:
            return None
        g = m.group(1).strip()
        return None if (g == "_No active goal._" or not g) else g
    except Exception:
        return None


def load_memory_into_context(path, label):
    if not path or not os.path.exists(path):
        return None
    try:
        content = open(path, "r", encoding="utf-8").read().strip()
        return f"[JARVIS {label.upper()} MEMORY]\n{content}" if content else None
    except Exception:
        return None


def python_trigger_memory_update(turn_tool_outputs, assistant_reply):
    combined = " ".join(turn_tool_outputs).lower() + " " + assistant_reply.lower()
    update_memory("session", f"Turn summary: {assistant_reply.strip()[:200]}")

    for p in re.findall(r'[a-zA-Z]:\\[^\s\'"<>|?*]+', assistant_reply):
        if any(ext in p.lower() for ext in [".py", ".md", ".json", ".txt", ".exe", ".ps1"]):
            update_memory("master", f"Referenced file path: {p}")
            break

    m = re.search(r"stdout:\s*\n(.+)", combined)
    if m:
        update_memory("session", f"Terminal output: {m.group(1).strip()[:120]}")

    if _active_project_memory_path and "success:" in combined:
        if re.search(r"[a-zA-Z]:\\[^\s]+", combined):
            update_memory("project", f"Action completed: {assistant_reply.strip()[:150]}")

    completion_signals = ["done", "completed", "finished", "task complete", "all done"]
    if _current_goal and any(s in combined for s in completion_signals):
        set_current_goal("none", reason="Python auto-detected completion")


def _bootstrap_all_files():
    """
    Create every folder and file Jarvis needs on first run.
    All calls are no-ops if the file already exists.
    Ship only main.py + gui.py — everything else is generated here.
    """
    # ── Directories ───────────────────────────────────────────────────────────
    for d in (TARGET_DIR, STORAGE_DIR, SKILLS_DIR):
        os.makedirs(d, exist_ok=True)

    # app_maps directory (used by UIA blueprint cache)
    os.makedirs(os.path.join(STORAGE_DIR, "app_maps"), exist_ok=True)

    # ── Helper: create a file only if it doesn't exist ────────────────────────
    def seed(path, content):
        if not os.path.exists(path):
            try:
                parent = os.path.dirname(path)
                if parent:
                    os.makedirs(parent, exist_ok=True)
                with open(path, "w", encoding="utf-8") as f:
                    f.write(content)
            except Exception as e:
                print(f"⚠️  Could not create {path}: {e}")

    ts = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")

    # ── Core knowledge files ──────────────────────────────────────────────────
    seed(COMMANDS_FILE,
        "# Jarvis Commands\n"
        "# Add preferred PowerShell commands below.\n"
        "# Format: `CommandName` — description\n\n"
        "## Commands\n"
        "- `Get-Location` — print the current working directory\n"
        "- `Get-ChildItem` — list folder contents\n"
        "- `Start-Process` — launch an application\n"
        "- `start` — shorthand to open files/apps\n"
    )

    seed(INSTRUCTIONS_FILE,
        "# Jarvis Instructions & Preferences\n"
        "# Add user preferences and behavioural rules below.\n"
        "# Format: one rule per line, starting with '- '\n\n"
        "## Preferences\n"
        "- Always use PowerShell commands, never CMD or Linux commands.\n"
        "- Use 'start' instead of 'Start-Process' when launching apps.\n"
    )

    seed(PATHS_FILE,
        "# Jarvis Paths\n"
        "# Absolute paths to applications, folders and files on this machine.\n\n"
        "## Paths\n"
    )

    seed(DOMAIN_INDEX,
        "# Jarvis Domain Knowledge Index\n"
        "# Registered domain-specific knowledge files.\n"
        "# Format: `filename_without_ext` - description\n\n"
        "## Files\n"
    )

    seed(DOMAIN_SKILLS_INDEX,
        "# Jarvis Domain Skills Index\n"
        "# Registered domain-specific skill files.\n"
        "# Format: [domain] `filename_without_ext` - description\n\n"
        "## Skills\n"
    )

    seed(SKILLS_INDEX,
        "# Jarvis Skills Index\n\n"
        "Each entry lists a skill filename and its description.\n"
        f"Skill files live in: {SKILLS_DIR}\n\n"
        "## Skills\n"
        "_No skills registered yet._\n"
    )

    # ── Memory files ──────────────────────────────────────────────────────────
    seed(MASTER_MEMORY,
        "# Jarvis Master Memory\n"
        f"Initialised: {ts}\n"
    )

    seed(SESSION_MEMORY,
        "# Jarvis Session Memory\n"
        f"Session started: {ts}\n\n"
        "## Current Goal\n"
        "_No active goal._\n\n"
        "## Goal History\n"
    )

    seed(RESPONSE_MEMORY, "")   # starts empty every run; cleared by set_current_goal(none)

    print("✅ [All Jarvis files and folders verified/created.]")


def init_memory_at_startup():
    global _active_project_memory_path, _current_goal
    _bootstrap_all_files()   # always runs first — safe no-op on subsequent launches
    os.makedirs(STORAGE_DIR, exist_ok=True)
    os.makedirs(SKILLS_DIR, exist_ok=True)
    _ensure_kb_files()
    injections = []

    master_ctx = load_memory_into_context(MASTER_MEMORY, "master")
    if master_ctx:
        injections.append(master_ctx)
        print("🧠 [Master memory loaded.]")
    else:
        print("🧠 [No master memory — starting fresh.]")
        if not os.path.exists(MASTER_MEMORY):
            write_local_file(MASTER_MEMORY, "# Jarvis Master Memory\n")

    if os.path.exists(SESSION_MEMORY):
        session_ctx = load_memory_into_context(SESSION_MEMORY, "session (continued)")
        if session_ctx:
            injections.append(session_ctx)
            print("🧠 [Session memory loaded.]")
        restored = get_current_goal_from_file()
        if restored:
            _current_goal = restored
            print(f"🎯 [Goal restored: {restored}]")
    else:
        ts = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
        write_local_file(
            SESSION_MEMORY,
            f"# Jarvis Session Memory\nSession started: {ts}\n\n"
            f"{GOAL_SECTION_HEADER}\n_No active goal._\n\n{GOAL_SECTION_END}\n"
        )
        print("🧠 [New session memory created.]")

    if os.path.exists(INSTRUCTIONS_FILE):
        try:
            with open(INSTRUCTIONS_FILE, "r", encoding="utf-8") as _f:
                _instr = _f.read().strip()
            if _instr:
                injections.append("[JARVIS INSTRUCTIONS — always active]\n" + _instr)
                print("📌 [Instructions loaded.]")
        except Exception:
            pass

    skill_count = len([f for f in os.listdir(SKILLS_DIR) if f.endswith(".md")])
    print(f"📋 [Skills: {skill_count} available]")

    print("\n──────────────────────────────────────────")
    project_name = input("📁 Active project (Enter to skip): ").strip()
    print("──────────────────────────────────────────")

    if project_name:
        project_dir  = os.path.join(r"D:\\", project_name)
        project_file = os.path.join(project_dir, "project_memory.md")
        _active_project_memory_path = project_file
        if os.path.exists(project_file):
            ctx = load_memory_into_context(project_file, f"project ({project_name})")
            if ctx:
                injections.append(ctx)
                print(f"🧠 [Project memory loaded: {project_file}]")
        else:
            os.makedirs(project_dir, exist_ok=True)
            write_local_file(
                project_file,
                f"# Project Memory: {project_name}\n"
                f"Created: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}\n"
            )
            print(f"🧠 [New project memory: {project_file}]")
        update_memory("master", f"Active project: {project_name} ({project_dir})")
    else:
        print("🧠 [No active project.]")

    return injections


# =============================================================================
# 7. PATH RESOLVER
# =============================================================================

_SYSTEM_FILES = {
    os.path.normcase(COMMANDS_FILE),
    os.path.normcase(INSTRUCTIONS_FILE),
    os.path.normcase(PATHS_FILE),
    os.path.normcase(DOMAIN_INDEX),
    os.path.normcase(DOMAIN_SKILLS_INDEX),
    os.path.normcase(MASTER_MEMORY),
    os.path.normcase(SESSION_MEMORY),
    os.path.normcase(RESPONSE_MEMORY),
    os.path.normcase(SKILLS_INDEX),
}


def _is_absolute(path):
    return os.path.isabs(path) or (len(path) > 1 and path[1] == ":")


def resolve_file_path(path):
    """
    Resolve a relative path using safe BFS with plain Get-ChildItem at each level.
    Never uses -Recurse or -Filter. Stops as soon as the filename is matched.
    """
    MAX_EXPLORE_DEPTH = 4

    if not path:
        return path, ""
    if _is_absolute(path):
        return path, ""

    filename = os.path.basename(path)
    for sp in _SYSTEM_FILES:
        if os.path.basename(sp) == os.path.normcase(filename):
            return sp, ""

    print(f"   [Resolver] '{path}' is relative — BFS exploring under {STARTUP_DIR}...")

    import platform
    from collections import deque

    def _list_entries(dirpath):
        try:
            if platform.system() == "Windows":
                ps = (
                    f"Get-ChildItem -Path '{dirpath}' | "
                    f"Select-Object Name,"
                    f"@{{n='D';e={{if($_.PSIsContainer){{'1'}}else{{'0'}}}}}} | "
                    f"ConvertTo-Csv -NoTypeInformation | Out-String"
                )
                out    = execute_terminal_command(ps)
                stdout = out.split("STDOUT:")[-1].split("STDERR:")[0].strip()
                entries = []
                for line in stdout.splitlines()[1:]:
                    line = line.strip().strip('"')
                    if not line:
                        continue
                    parts = [p.strip().strip('"') for p in line.split('","')]
                    if len(parts) >= 2:
                        name, is_dir = parts[0], parts[1] == "1"
                        entries.append((name, is_dir, os.path.join(dirpath, name)))
                return entries
            else:
                return [
                    (e, os.path.isdir(os.path.join(dirpath, e)), os.path.join(dirpath, e))
                    for e in os.listdir(dirpath)
                ]
        except Exception:
            return []

    queue   = deque([(STARTUP_DIR, 0)])
    visited = set()
    while queue:
        current_dir, depth = queue.popleft()
        if current_dir in visited or depth > MAX_EXPLORE_DEPTH:
            continue
        visited.add(current_dir)
        for name, is_dir, full_path in _list_entries(current_dir):
            if name.lower() == filename.lower() and not is_dir:
                msg = f"Resolved '{path}' -> '{full_path}'"
                print(f"   [Resolver] {msg}")
                return full_path, msg
            if is_dir and depth < MAX_EXPLORE_DEPTH:
                queue.append((full_path, depth + 1))

    msg = f"Could not find '{filename}' within {MAX_EXPLORE_DEPTH} levels of {STARTUP_DIR}. Path used as given."
    print(f"   [Resolver] {msg}")
    return path, msg

# =============================================================================
# 8. COMMAND SAFETY — advisory lookup
# =============================================================================

def _load_commands_whitelist():
    if not os.path.exists(COMMANDS_FILE):
        return set()
    try:
        content = open(COMMANDS_FILE, "r", encoding="utf-8").read()
        tokens  = re.findall(r"`([^`]+)`", content)
        for line in content.splitlines():
            stripped = line.strip().lstrip("-#* ")
            if stripped:
                tokens.append(stripped.split()[0])
        return {t.lower() for t in tokens if t.strip()}
    except Exception:
        return set()


def _command_looks_known(cmd, whitelist):
    if not whitelist:
        return True
    first = cmd.strip().split()[0].lower() if cmd.strip() else ""
    return first in whitelist


# =============================================================================
# 9. PERSISTENT ORCHESTRATION ENGINE
# =============================================================================

# Words that signal a trivial/short turn — skip Gemini pre-reasoning for these
_TRIVIAL_PATTERNS = {
    "yes","no","ok","okay","y","n","sure","fine","good","thanks","thank you",
    "exit","quit","new session","stop","cancel","abort","go ahead","run it",
    "grant","approve","continue","done","next","skip","wait","hello","hi",
    "hey","what","why","how","who","when","where",
}

def _is_trivial_input(text: str) -> bool:
    """Return True if the input is short/simple enough to skip Gemini pre-reasoning."""
    stripped = text.strip().lower()
    # Single word or very short
    if len(stripped.split()) <= 2:
        return True
    # Matches a known trivial phrase
    if stripped in _TRIVIAL_PATTERNS:
        return True
    # Pure approval/bypass turns
    if "[USER MANUALLY GRANTED BYPASS]" in text:
        return True
    return False


def get_gemini_reasoning(user_input: str, conversation_history: list) -> str | None:
    """
    Ask Gemini to reason about the user request and produce a plan/analysis
    that Jarvis will use as context before deciding on tool calls.

    Gemini receives:
      - A compact summary of the last few turns for context
      - The user's current request
      - A prompt asking for a structured plan or answer

    Returns the reasoning text, or None if Gemini is unavailable.
    """
    if not _GEMINI_AVAILABLE or _gemini_client is None:
        return None
    try:
        # Build a compact conversation summary for context (last 6 non-system messages)
        non_sys = [m for m in conversation_history if m.get("role") != "system"]
        recent  = non_sys[-6:] if len(non_sys) > 6 else non_sys
        history_text = ""
        for m in recent:
            role    = m.get("role", "?")
            content = m.get("content", "")
            if isinstance(content, str):
                history_text += f"{role.upper()}: {content[:300]}\n"

        prompt = (
            "You are the reasoning brain of Jarvis, a Windows desktop AI agent. "
            "Jarvis (a local model) will execute tool calls based on your plan.\n\n"
            f"RECENT CONVERSATION:\n{history_text}\n"
            f"CURRENT USER REQUEST: {user_input}\n\n"
            "Your job:\n"
            "1. Analyse the request carefully.\n"
            "2. Identify any ambiguities, risks, or things Jarvis should check first.\n"
            "3. Write a clear, ordered action plan for Jarvis to follow.\n"
            "4. If the request is a question, provide your best answer directly.\n"
            "5. Flag anything Jarvis should NOT do or should be careful about.\n\n"
            "Be concise. Jarvis will read your output before acting."
        )

        # Pick a credit-safe model (Defaults to 2.0-flash or safe 2.5-flash)
        model_id, rationale = _pick_gemini_model(user_input)
        
        # Modern unified API request syntax using the global client
        response = _gemini_client.models.generate_content(
            model=model_id,
            contents=prompt,
        )
        reasoning = response.text.strip()
        
        print(f"\n🤖 [Gemini/{model_id} reasoning: {len(reasoning)} chars]")
        return reasoning
    except Exception as e:
        print(f"\n⚠️  [Gemini reasoning failed: {e}]")
        return None

# Tools that count as "verification" — capped so Jarvis can't loop forever
_VERIFY_TOOLS = {"fallback_view_screen", "fallback_find_text"}
# Maximum consecutive verification tool calls allowed before we force a reply
MAX_VERIFY_CALLS = 2

def wait(seconds: float) -> str:
    """Pauses thread execution for the specified duration."""
    try:
        time.sleep(seconds)
        return f"Successfully paused for {seconds} seconds."
    except Exception as e:
        return f"Error during wait execution: {str(e)}"

# =============================================================================
# LEGACY TOOL-CALL FALLBACK (qwen2.5-coder and similar)
# =============================================================================
# Models in LEGACY_TOOLCALL_MODELS are still sent the exact same `tools=`
# schema as modern models — Ollama still injects it into their chat template.
# The difference is purely on the READ side: these models frequently put the
# tool call as plain text inside `content` (raw JSON, or wrapped in
# <tool_call></tool_call> tags per the Qwen2.5 template) instead of Ollama's
# structured `tool_calls` field. This block of code ONLY runs as a fallback
# when `response["message"].get("tool_calls")` is already empty, so it never
# touches or alters behavior for models that report tool_calls natively.

# Matches a <tool_call> ... </tool_call> block (Qwen2.5 chat template style)
_TOOLCALL_TAG_RE = re.compile(r"<tool_call>\s*(\{.*?\})\s*</tool_call>", re.DOTALL)
# Matches a ```json fenced block
_JSON_FENCE_RE   = re.compile(r"```(?:json)?\s*(\{.*?\})\s*```", re.DOTALL)

def _find_balanced_json_objects(text: str):
    """
    Scan text for top-level {...} objects using brace counting (not regex),
    so nested objects like {"name": "x", "arguments": {"a": 1}} are captured
    whole instead of being cut off at the first inner '}'. Returns a list of
    (start, end, blob) tuples for every balanced top-level object found.
    """
    results = []
    depth = 0
    start = None
    in_string = False
    escape = False
    for i, ch in enumerate(text):
        if in_string:
            if escape:
                escape = False
            elif ch == "\\":
                escape = True
            elif ch == '"':
                in_string = False
            continue
        if ch == '"':
            in_string = True
        elif ch == "{":
            if depth == 0:
                start = i
            depth += 1
        elif ch == "}":
            if depth > 0:
                depth -= 1
                if depth == 0 and start is not None:
                    results.append((start, i + 1, text[start:i + 1]))
                    start = None
    return results

_KNOWN_TOOL_NAMES = None  # lazily populated from `tools` schema, see below

def _known_tool_names():
    global _KNOWN_TOOL_NAMES
    if _KNOWN_TOOL_NAMES is None:
        _KNOWN_TOOL_NAMES = {t["function"]["name"] for t in tools}
    return _KNOWN_TOOL_NAMES

def _try_parse_tool_json(blob: str):
    """Parse a JSON blob into a normalized tool_call dict, or None if invalid."""
    try:
        obj = json.loads(blob)
    except Exception:
        return None
    if not isinstance(obj, dict):
        return None
    name = obj.get("name") or obj.get("function") or obj.get("tool")
    args = obj.get("arguments", obj.get("parameters", {}))
    if not name or name not in _known_tool_names():
        return None
    if isinstance(args, str):
        try:
            args = json.loads(args)
        except Exception:
            pass
    return {"function": {"name": name, "arguments": args}}

def _extract_legacy_tool_calls(content: str):
    """
    Scan a model's free-text `content` for one or more tool calls when the
    structured tool_calls field came back empty. Returns (tool_calls, leftover_text)
    where leftover_text is the content with the recognized tool-call JSON stripped
    out (so it isn't shown to the user / re-fed as a duplicate plain message).
    Returns ([], content) if nothing parseable is found — caller treats that
    as a normal plain-text reply, completely transparent for modern models.
    """
    if not content or not content.strip():
        return [], content

    found = []
    cleaned = content

    # 1) <tool_call>...</tool_call> tags — may be one or several
    tag_matches = list(_TOOLCALL_TAG_RE.finditer(content))
    if tag_matches:
        for m in tag_matches:
            parsed = _try_parse_tool_json(m.group(1))
            if parsed:
                found.append(parsed)
        if found:
            cleaned = _TOOLCALL_TAG_RE.sub("", content).strip()
            return found, cleaned

    # 2) ```json fenced block
    fence_match = _JSON_FENCE_RE.search(content)
    if fence_match:
        parsed = _try_parse_tool_json(fence_match.group(1))
        if parsed:
            cleaned = _JSON_FENCE_RE.sub("", content, count=1).strip()
            return [parsed], cleaned

    # 3) Bare JSON object(s) anywhere in the text (most common qwen2.5-coder
    #    case: the ENTIRE content is just the JSON object, nothing else).
    #    Brace-balanced scan handles nested "arguments": {...} correctly.
    candidates = _find_balanced_json_objects(content)
    if candidates:
        consumed_spans = []
        for start, end, blob in candidates:
            parsed = _try_parse_tool_json(blob)
            if parsed:
                found.append(parsed)
                consumed_spans.append((start, end))
        if found:
            cleaned = content
            for start, end in sorted(consumed_spans, reverse=True):
                cleaned = cleaned[:start] + cleaned[end:]
            return found, cleaned.strip()

    return [], content


def _call_ollama(messages, result_q):
    """Run ollama.chat on a background thread and put the result in result_q."""
    try:
        resp = ollama.chat(model=MODEL_NAME, messages=messages, tools=tools)

        # ── Legacy fallback: only engages if native tool_calls is empty AND
        #    the active model is a known weak-tool-calling family. Modern
        #    models always have tool_calls populated natively and never reach
        #    this branch, so their behavior/performance is unchanged. ──────────
        if not resp["message"].get("tool_calls") and _is_legacy_toolcall_model(MODEL_NAME):
            raw_content = resp["message"].get("content") or ""
            legacy_calls, cleaned_content = _extract_legacy_tool_calls(raw_content)
            if legacy_calls:
                # Normalize to a plain dict before mutating — the ollama
                # client's Message is a pydantic SubscriptableBaseModel and
                # may not support item assignment the way a plain dict does.
                # Converting here keeps `resp["message"]` behaving exactly
                # like the plain-dict messages used everywhere else in the
                # conversation history (e.g. tool/user messages appended
                # manually elsewhere in this file).
                msg_dict = dict(resp["message"])
                msg_dict["tool_calls"] = legacy_calls
                msg_dict["content"]    = cleaned_content
                msg_dict.setdefault("role", "assistant")
                try:
                    resp["message"] = msg_dict
                except (TypeError, KeyError):
                    # ChatResponse may not support item assignment either —
                    # fall back to attribute assignment, then to wrapping the
                    # whole response in a plain dict as a last resort so the
                    # rest of the pipeline always gets dict-like access.
                    try:
                        resp.message = msg_dict
                    except Exception:
                        resp = {**dict(resp), "message": msg_dict}

        result_q.put(("ok", resp))
    except Exception as e:
        result_q.put(("err", e))


MAX_ACTION_TRIES = 3

def get_system_prompt():
    """Generates the master system prompt dynamically based on loaded modules."""
    ocr_status = (
        "fallback_click_text(text) uses Tesseract OCR to find and click any visible "
        "text element precisely — no coordinate estimation needed. "
        "fallback_find_text(text) returns coordinates without clicking. "
        "Prefer these over fallback_view_screen + fallback_click_grid for any element with readable text."
        if _TESSERACT_AVAILABLE else
        "Tesseract OCR is not installed — fallback_click_text and fallback_find_text "
        "are unavailable. Use fallback_view_screen + fallback_click_grid with grid coordinates instead."
    )

    # Only injected for models known to be unreliable at native structured
    # tool_calls (see LEGACY_TOOLCALL_MODELS / _is_legacy_toolcall_model).
    # Modern models never see this text — get_system_prompt() output is
    # otherwise byte-for-byte identical to before this change.
    legacy_toolcall_hint = (
        "\n- TOOL CALLS: when you need to use a tool, output ONLY a single JSON "
        "object — nothing else, no extra commentary before or after it — in "
        "exactly this shape: {\"name\": \"<tool_name>\", \"arguments\": {<args>}}. "
        "Wrapping it in <tool_call></tool_call> tags is also fine. Never "
        "describe the tool call in prose instead of emitting this JSON."
        if _is_legacy_toolcall_model(MODEL_NAME) else ""
    )

    return (
        "You are Jarvis, a Windows desktop AI agent. Rules:\n"
        "- NEVER fabricate outputs. Always use tools.\n"
        "- PowerShell only. No CMD, no Linux/bash.\n"
        f"- Commands: run directly; check commands.md if unsure "
        f"({COMMANDS_FILE}); search internet as last resort.\n"
        "- PATH NAVIGATION: never assume a path exists. "
        "Check paths.md first. If not there, call explore_path(parent_dir) "
        "to list what is actually present, then drill down one level at a time.\n"
        f"- RETRY CAP: max {MAX_ACTION_TRIES} attempts per action. "
        "If exceeded, stop and report to the user.\n"
        "- Instructions: your rules are loaded at startup above. "
        "Follow them always. Add new ones with add_instruction.\n"
        f"- Skills: list_skills then load_skill for multi-step tasks. "
        f"Dir: {SKILLS_DIR}\n"
        "- GUI INTERACTION — DEFAULT to UI Automation. ONLY use OCR/screen-viewing if "
        "the user explicitly asks for it, OR click_ui_element says the window is "
        "canvas/WebGL-rendered (UIA literally cannot see inside it):\n"
        "  DEFAULT: click_ui_element(window_title, description) — ONE call for ANY "
        "desktop app interaction (VS Code, Explorer, Claude, Discord, settings, browsers, "
        "etc). Just describe the element in plain English, e.g. "
        "click_ui_element('Claude', 'Close button') or "
        "click_ui_element('Notepad', 'Save'). It searches the whole window, clicks the "
        "best match, and automatically falls back to clicking the element's on-screen "
        "position if it can't be invoked directly (handles Electron app window controls "
        "correctly). If you get 'No element matched', it lists the actual element names "
        "available — retry with one of those exact names. Always use this unless one of "
        "the exceptions below applies.\n"
        f"  EXCEPTION — OCR (only if user explicitly asks, or click_ui_element reports "
        "canvas/WebGL rendering): {ocr_status}\n"
        f"  EXCEPTION — GRID (only if user explicitly asks for screen viewing/coordinates, "
        f"or both above fail): fallback_view_screen returns a "
        f"{MODEL_CANVAS_W}x{MODEL_CANVAS_H} canvas with yellow grid labels. "
        f"Read the label nearest the target and pass it to fallback_click_grid. "
        f"Python multiplies by ({SCALE_X:.2f}, {SCALE_Y:.2f}) automatically.\n"
        "  manual_scan_app_layouts/manual_inspect_app_subtree/manual_interact_with_ui are manual/advanced "
        "tools — only use them if click_ui_element explicitly fails and you need to "
        "explore available elements yourself.\n"
        "  Never use fallback_view_screen, fallback_find_text, or fallback_click_text "
        "unprompted — click_ui_element is the default for everything GUI-related.\n"
        "- After a click, verify ONLY if you are genuinely uncertain the action succeeded. Do not verify routinely — trust the tool result unless it reported an error.\n"
        "- Memory: update_memory(master|project|session, content).\n"
        "- Goals: set_current_goal on task start; goal='none' on completion.\n"
        f"- paths.md ({PATHS_FILE}): read when you need a path. Add with add_path.\n"
        "- Domain knowledge: list_domain_knowledge, read_domain_knowledge, create_domain_knowledge.\n"
        "- Domain skills: list_domain_skills, load_skill, create_domain_skill.\n"
        "- Gemini is your reasoning teammate. You have a tool called consult_gemini. "
        "Call it explicitly when you face complex coding tasks, architecture planning, "
        "or need a deep analysis that exceeds your local context.\n"
        "- FILE FORMATS: use read_file_smart for .pdf .docx .txt and all text formats. "
        "PDF/DOCX are read-only. Write with write_local_file (.txt/.md).\n"
        "- RESPONSE MEMORY (scratchpad, wiped each turn): for ANY multi-step task or "
        "document with multiple instructions, FIRST call write_response_memory with a "
        "numbered checklist, THEN work through each item calling append_response_memory "
        "to log progress. For large documents: chunk plan in response memory, process "
        "each chunk, append partial results, read_response_memory, then assemble final answer.\n"
        "- ANTI-HALLUCINATION: state what you are about to do before each action. "
        "If uncertain about a fact or path, verify with a tool — never guess.\n"
        "- Replies: Markdown."
        f"{legacy_toolcall_hint}"
    )



def process_chat_turn(conversation_history):
    clear_response_memory()
    turn_tool_outputs  = []
    whitelist          = _load_commands_whitelist()
    verify_call_count  = 0
    action_attempt_counts: dict = {}
    _abort_event.clear()

    # Keep system messages always; slide a window over the rest
    HISTORY_WINDOW = 20

    while True:
        # ── Ctrl+Q abort check ────────────────────────────────────────────────
        if _abort_event.is_set():
            print("\n🛑 [Response aborted by Ctrl+Q]")
            return "[Response terminated by user.]", turn_tool_outputs

        sys_msgs = [m for m in conversation_history if m.get("role") == "system"]
        non_sys  = [m for m in conversation_history if m.get("role") != "system"]
        trimmed  = sys_msgs + non_sys[-HISTORY_WINDOW:]

        # ── Run ollama.chat on a thread so Ctrl+Q can interrupt the wait ──────
        result_q = _queue.Queue()
        t = threading.Thread(target=_call_ollama, args=(trimmed, result_q), daemon=True)
        t.start()
        while t.is_alive():
            if _abort_event.is_set():
                print("\n🛑 [Response aborted by Ctrl+Q]")
                return "[Response terminated by user.]", turn_tool_outputs
            t.join(timeout=0.1)   # check abort flag every 100 ms

        status, payload = result_q.get()
        if status == "err":
            return f"[Ollama error: {payload}]", turn_tool_outputs
        response   = payload
        tool_calls = response["message"].get("tool_calls")

        if not tool_calls:
            conversation_history.append(response["message"])
            return response["message"]["content"] or "", turn_tool_outputs

        # ── Verification loop cap ─────────────────────────────────────────────
        all_verify = all(
            tc["function"]["name"] in _VERIFY_TOOLS for tc in tool_calls
        )
        if all_verify:
            verify_call_count += 1
            if verify_call_count > MAX_VERIFY_CALLS:
                # Force the model to stop verifying and give a final reply
                conversation_history.append(response["message"])
                conversation_history.append({
                    "role": "user",
                    "content": (
                        "[SYSTEM]: You have verified the result enough times. "
                        "Stop calling fallback_view_screen or fallback_find_text. "
                        "Give your final plain-text reply to the user now."
                    )
                })
                continue
        else:
            verify_call_count = 0   # reset counter when a real action runs

        conversation_history.append(response["message"])
        print(f"\n⚡ Jarvis requested {len(tool_calls)} action(s)...")

        needs_lookup = False
        unknown_cmd  = ""

        for tool in tool_calls:
            func_name = tool["function"]["name"]
            raw_args  = tool["function"]["arguments"]
            if isinstance(raw_args, str):
                try:    arguments = json.loads(raw_args)
                except: arguments = {}
            else:
                arguments = raw_args

            print(f" -> Executing: '{func_name}'")
            tool_images = None

            # ── Hard retry cap ────────────────────────────────────────────────
            _EXEMPT = {"update_memory","set_current_goal","add_instruction","add_path","explore_path","write_response_memory","append_response_memory","read_response_memory"}
            _cap_hit = False
            if func_name not in _EXEMPT:
                _karg = next((str(arguments[k])[:80] for k in
                    ("command","path","text","query","prompt","skill_name","name","instruction")
                    if k in arguments), "")
                _akey = (func_name, _karg)
                action_attempt_counts[_akey] = action_attempt_counts.get(_akey, 0) + 1
                if action_attempt_counts[_akey] > MAX_ACTION_TRIES:
                    cap_msg = (f"[RETRY CAP] '{func_name}' attempted {MAX_ACTION_TRIES} "
                               f"times with the same argument and has not succeeded. "
                               f"Stop retrying immediately. Tell the user what failed "
                               f"and ask how they want to proceed.")
                    print(f"\n🚫 [Retry cap reached for '{func_name}']")
                    turn_tool_outputs.append(cap_msg)
                    conversation_history.append({"role":"tool","content": cap_msg})
                    _cap_hit = True
            if _cap_hit:
                break   # exits the for-tool loop; then the while loop gets
                        # one more model call to produce the failure reply

            # ── Tool dispatch ──────────────────────────────────────────────────
            if func_name == "read_local_file":
                raw_path           = arguments.get("path", "")
                resolved, res_msg  = resolve_file_path(raw_path)
                file_result        = read_local_file(resolved)
                tool_output        = (f"[PATH RESOLVED: {res_msg}]\n{file_result}"
                                      if res_msg else file_result)

            elif func_name == "write_local_file":
                raw_path          = arguments.get("path", "")
                resolved, res_msg = resolve_file_path(raw_path)
                tool_output       = write_local_file(resolved, arguments.get("content"))
                if res_msg:
                    tool_output = f"[PATH RESOLVED: {res_msg}] {tool_output}"

            elif func_name == "append_local_file":
                raw_path          = arguments.get("path", "")
                resolved, res_msg = resolve_file_path(raw_path)
                tool_output       = append_local_file(resolved, arguments.get("content"))
                if res_msg:
                    tool_output = f"[PATH RESOLVED: {res_msg}] {tool_output}"

            elif func_name == "search_internet":
                tool_output = search_internet(arguments.get("query"))

            elif func_name == "execute_terminal_command":
                cmd = arguments.get("command", "").strip()
                if not cmd:
                    tool_output = "Error: No command provided."
                else:
                    print(f"   [Terminal] > {cmd}")
                    if whitelist and not _command_looks_known(cmd, whitelist):
                        needs_lookup = True
                        unknown_cmd  = cmd
                        print(f"   [ℹ️  '{cmd.split()[0]}' not in commands.md]")
                    tool_output = execute_terminal_command(
                        cmd, working_directory=arguments.get("working_directory")
                    )

            elif func_name == "fallback_view_screen":
                b64_img = capture_screen_to_ram()
                if not b64_img.startswith("Error"):
                    tool_output = (
                        f"Screenshot captured at canvas size {MODEL_CANVAS_W}x{MODEL_CANVAS_H}. "
                        f"Yellow grid labels are canvas coordinates. "
                        f"Pass them directly to fallback_click_grid — Python scales by "
                        f"({SCALE_X:.2f}x, {SCALE_Y:.2f}x) to reach real screen pixels. "
                        f"For text elements, prefer fallback_click_text for precision."
                    )
                    tool_images = [b64_img]
                else:
                    tool_output = b64_img

            elif func_name == "fallback_find_text":
                tool_output = fallback_find_text(arguments.get("text", ""))

            elif func_name == "fallback_click_grid":
                x          = arguments.get("x", 0)
                y          = arguments.get("y", 0)
                click_type = arguments.get("click_type", "left_click")
                print(f"   [Click] {click_type} at canvas ({x},{y})")
                tool_output = fallback_click_grid(x, y, click_type)

            elif func_name == "fallback_click_text":
                text       = arguments.get("text", "")
                click_type = arguments.get("click_type", "left_click")
                print(f"   [OCR Click] '{text}'")
                tool_output = fallback_click_text(text, click_type)

            elif func_name == "type_text":
                text        = arguments.get("text", "")
                special_key = arguments.get("special_key", None)
                print(f"   [Type] '{text[:40]}{'...' if len(text)>40 else ''}'")
                tool_output = type_text(text, special_key)

            elif func_name == "update_memory":
                tool_output = update_memory(
                    arguments.get("target", "session"),
                    arguments.get("content", "")
                )

            elif func_name == "set_current_goal":
                tool_output = set_current_goal(
                    arguments.get("goal", ""),
                    arguments.get("reason", "")
                )

            elif func_name == "list_skills":
                tool_output = list_skills()

            elif func_name == "load_skill":
                tool_output = load_skill(arguments.get("skill_name", ""))

            elif func_name == "read_instructions":
                tool_output = read_instructions()

            elif func_name == "add_instruction":
                tool_output = add_instruction(arguments.get("instruction", ""))

            elif func_name == "read_paths":
                tool_output = read_paths()

            elif func_name == "explore_path":
                tool_output = explore_path(arguments.get("path", STARTUP_DIR))

            elif func_name == "add_path":
                tool_output = add_path(
                    arguments.get("label", ""),
                    arguments.get("path", ""),
                    arguments.get("note", "")
                )

            elif func_name == "create_domain_knowledge":
                tool_output = create_domain_knowledge(
                    arguments.get("name", ""),
                    arguments.get("description", ""),
                    arguments.get("initial_content", "")
                )

            elif func_name == "list_domain_knowledge":
                tool_output = list_domain_knowledge()

            elif func_name == "read_domain_knowledge":
                tool_output = read_domain_knowledge(arguments.get("name", ""))

            elif func_name == "create_domain_skill":
                tool_output = create_domain_skill(
                    arguments.get("name", ""),
                    arguments.get("domain", ""),
                    arguments.get("description", ""),
                    arguments.get("content", "")
                )

            elif func_name == "list_domain_skills":
                tool_output = list_domain_skills()

            elif func_name == "consult_gemini":
                tool_output = consult_gemini(
                    arguments.get("prompt", ""),
                    arguments.get("task_type", "auto"),
                    arguments.get("context", "")
                )
            elif func_name == "read_file_smart":
                tool_output = read_file_smart(arguments.get("path", ""))
            elif func_name == "read_file_chunk":
                tool_output = read_file_chunk(arguments.get("path",""), int(arguments.get("chunk_index",1)))
            elif func_name == "write_docx_file":
                tool_output = write_docx_file(arguments.get("path",""), arguments.get("content",""))
            elif func_name == "write_response_memory":
                tool_output = write_response_memory(arguments.get("content",""))
            elif func_name == "append_response_memory":
                tool_output = append_response_memory(arguments.get("content",""))
            elif func_name == "read_response_memory":
                tool_output = read_response_memory()
            elif func_name == "manual_scan_app_layouts":
                tool_output = manual_scan_app_layouts(arguments.get("window_title", ""))

            elif func_name == "manual_inspect_app_subtree":
                tool_output = manual_inspect_app_subtree(
                    arguments.get("window_title", ""),
                    arguments.get("subtree_key", "")
                )

            elif func_name == "click_ui_element":
                window_title = arguments.get("window_title", "")
                description  = arguments.get("description", "")
                action       = arguments.get("action", "click")
                text_to_type = arguments.get("text_to_type", "")
                print(f"   [UIA] click_ui_element: '{description}' in '{window_title}' (action={action})")
                tool_output = click_ui_element(window_title, description, action, text_to_type)
                if tool_output.startswith("Success"):
                    print(f"   [UIA] ✅ {tool_output}")
                else:
                    print(f"   [UIA] ⚠️  {tool_output[:120]}")

            elif func_name == "manual_interact_with_ui":
                print(f"   [UIA] {arguments.get('action')} on {arguments.get('property_value')}")
                tool_output = manual_interact_with_ui(
                    arguments.get("window_title", ""),
                    arguments.get("control_type", ""),
                    arguments.get("search_property", ""),
                    arguments.get("property_value", ""),
                    arguments.get("action", ""),
                    arguments.get("text_to_type", "")
                )

            elif func_name == "list_active_windows":
                tool_output = list_active_windows()

            elif func_name == "read_aggregated_text":
                print(f"   [UIA] Aggregating text blocks from: '{arguments.get('window_title')}'")
                if _UIA_AVAILABLE:
                    tool_output = ui_navigator.read_aggregated_text(
                        window_title=arguments.get("window_title", ""),
                        container_key=arguments.get("container_key", None)
                    )
                else:
                    tool_output = "UIA library not available."
            
            elif func_name == "query_gemini_app":
                prompt_payload = arguments.get("prompt", "")
                print(f"   [Bridge] Handing task execution off to Gemini Application...")
                if _UIA_AVAILABLE:
                    tool_output = ui_navigator.query_gemini_app(prompt=prompt_payload)
                else:
                    tool_output = "Execution failed: UIA layer is unavailable."
            
            elif func_name == "manage_gemini_chat":
                tool_output = ui_navigator.manage_gemini_chat(
                    action=arguments.get("action"),
                    chat_name=arguments.get("chat_name")
                )
            else:
                tool_output = f"Unknown tool: {func_name}"

            
            # ──────────────────────────────────────────────────────────────────

            tool_output = tool_output.replace("<", "&lt;").replace(">", "&gt;")
            turn_tool_outputs.append(tool_output)

            msg = {"role": "tool", "content": tool_output}
            if tool_images:
                msg["images"] = tool_images
            conversation_history.append(msg)

        if needs_lookup:
            conversation_history.append({
                "role": "user",
                "content": (
                    f"[SYSTEM NOTE]: '{unknown_cmd.split()[0]}' was not in commands.md. "
                    "Check commands.md or search online if you are unsure it was correct."
                )
            })


# =============================================================================
# 10. INTERACTIVE MAIN LOOP
# =============================================================================

if __name__ == "__main__":
    os.makedirs(TARGET_DIR, exist_ok=True)

    # ── Ctrl+Q abort hotkey ────────────────────────────────────────────────────
    if _KEYBOARD_AVAILABLE:
        _keyboard.add_hotkey("ctrl+q", lambda: _abort_event.set())
        print("⌨️  [Ctrl+Q registered — press to abort the current response]")
    else:
        print("⚠️  [keyboard package not found — Ctrl+Q abort unavailable]")
        print("    Install with: pip install keyboard")

    # ── Tesseract status ───────────────────────────────────────────────────────
    if _TESSERACT_AVAILABLE:
        print("👁️  [Tesseract OCR: available — fallback_click_text is active]")
    else:
        print("⚠️  [Tesseract OCR not found — fallback_click_text will not work]")
        print(f"    Install from: https://github.com/UB-Mannheim/tesseract/wiki")
        print(f"    Then: pip install pytesseract")
        print(f"    Verify TESSERACT_PATH = {TESSERACT_PATH}")

    # ── Gemini status ──────────────────────────────────────────────────────────
    if _GEMINI_AVAILABLE:
        print("🤖 [Gemini: available — consult_gemini is active]")
    else:
        print(f"⚠️  [Gemini not available: {_gemini_load_msg}]")
        print(f"    Secrets file expected at: {os.path.abspath(SECRETS_FILE)}")
        print( "    Format: { \"GEMINI_API_KEY\": \"your_key_here\" }")
        print( "    Install: pip install google-genai")

    memory_injections = init_memory_at_startup()

    try:
        with open(LOG_FILE, "w", encoding="utf-8") as f:
            f.write("# Jarvis Master Interaction Log\n")
            f.write(f"Session started: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}\n")
            f.write("=========================================\n\n")
    except Exception as e:
        print(f"⚠️ Warning: Could not initialise log file: {e}")

    print("\n====================================================")
    print("Jarvis local agent started. Persistent Chat Ready.")
    print(f"Tracking live session in: {LOG_FILE}")
    print("Type 'new session' to wipe session memory.")
    print("Type 'exit' or 'quit' to close.")
    print("====================================================\n")

    ocr_status = (
        "fallback_click_text(text) uses Tesseract OCR to find and click any visible "
        "text element precisely — no coordinate estimation needed. "
        "fallback_find_text(text) returns coordinates without clicking. "
        "Prefer these over fallback_view_screen + fallback_click_grid for any element with readable text."
        if _TESSERACT_AVAILABLE else
        "Tesseract OCR is not installed — fallback_click_text and fallback_find_text "
        "are unavailable. Use fallback_view_screen + fallback_click_grid with grid coordinates instead."
    )

    system_prompt = get_system_prompt()

    goal_reminder = ""
    if _current_goal:
        goal_reminder = (
            f"\n\n[GOAL REMINDER]\nCurrent goal: {_current_goal}\n"
            "Continue unless redirected."
        )

    history = [{"role": "system", "content": system_prompt + goal_reminder}]
    for inj in memory_injections:
        history.append({"role": "system", "content": inj})

    turn_counter = 1

    while True:
        try:
            user_input = input("\nYou: ").strip()

            if user_input.lower() == "new session":
                if os.path.exists(SESSION_MEMORY):
                    os.remove(SESSION_MEMORY)
                    print("🗑️ Session memory cleared.")
                _current_goal = None
                ts = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
                write_local_file(
                    SESSION_MEMORY,
                    f"# Jarvis Session Memory\nSession started: {ts}\n\n"
                    f"{GOAL_SECTION_HEADER}\n_No active goal._\n\n{GOAL_SECTION_END}\n"
                )
                print("🧠 New session started.\n")
                history = [{"role": "system", "content": system_prompt}]
                for inj in memory_injections[:1]:
                    history.append({"role": "system", "content": inj})
                turn_counter = 1
                continue

            if user_input.lower() in ("exit", "quit"):
                print("\nCleaning up...")
                if os.path.exists(LOG_FILE):
                    try:
                        os.remove(LOG_FILE)
                        print(f"🗑️ Deleted: {LOG_FILE}")
                    except Exception as e:
                        print(f"⚠️ Could not delete log: {e}")
                print("Goodbye!")
                break

            if not user_input:
                continue

            approval_keywords = ["yes", "grant", "approve", "run it", "go ahead", "y"]
            if any(kw in user_input.lower() for kw in approval_keywords):
                payload = f"{user_input} [USER MANUALLY GRANTED BYPASS]"
            else:
                payload = (
                    f"{user_input}\n\n"
                    "[SYSTEM]: For any multi-step task or document processing: "
                    "call write_response_memory with a numbered plan FIRST, then execute. "
                    "If this requires a shell command, call execute_terminal_command. "
                    "If it requires interacting with an app window, call click_ui_element "
                    "(UI Automation) — do not use screen/OCR tools unless explicitly asked."
                )

            history.append({"role": "user", "content": payload})

            print("\n[Thinking...]")
            assistant_reply, turn_tool_outputs = process_chat_turn(history)
            _print_reply("Jarvis:", assistant_reply)

            python_trigger_memory_update(turn_tool_outputs, assistant_reply)

            try:
                with open(LOG_FILE, "a", encoding="utf-8") as f:
                    f.write(f"# Response {turn_counter}\n\n")
                    f.write(f"### **User Prompt:**\n> {user_input}\n\n")
                    f.write(f"_Goal: {_current_goal or 'none'}_\n\n")
                    f.write(f"### **Jarvis Reply:**\n{assistant_reply}\n\n---\n\n")
                print(f"💾 [Logged response {turn_counter}]")
                turn_counter += 1
            except Exception as e:
                print(f"⚠️ Could not append to log: {e}")

        except KeyboardInterrupt:
            print("\n\nAborted.")
            if os.path.exists(LOG_FILE):
                try: os.remove(LOG_FILE)
                except: pass
            break
        except Exception as e:
            print(f"\nUnexpected error: {e}")